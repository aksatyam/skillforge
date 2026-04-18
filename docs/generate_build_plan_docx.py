#!/usr/bin/env python3
"""
Generate SkillForge AI Build Plan DOCX using the enterprise master-document
template standards from ~/.claude/CLAUDE.md.

Brand colors, status badges, and structure match existing Ashish-authored
artifacts (TPE master docs, VAPT reports).

Usage:
    python3 generate_build_plan_docx.py
    -> writes SkillForge_AI_Build_Plan.docx next to the source
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- Brand palette (from global CLAUDE.md Enterprise Document Standards) ---
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x27, 0xAE, 0x60)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
RED = RGBColor(0xE7, 0x4C, 0x3C)
DARK = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM = RGBColor(0x7F, 0x8C, 0x8D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SHADE_NAVY = "1B3A5C"
SHADE_ALT_ROW = "F8F9FA"
SHADE_LIGHT = "F0F4F8"
SHADE_SUCCESS = "E8F5E9"
SHADE_WARNING = "FFF8E1"
SHADE_DANGER = "FFEBEE"

DOC_ID = "SF-BUILD-PLAN-2026-001"
DOC_DATE = date(2026, 4, 18).isoformat()
AUTHOR = "Ashish Kumar Satyam"
COMPANY = "TechDigital WishTree"
FOR_ORG = "Qualtech"
VERSION = "v1.0"


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_font(cell, *, bold=False, color=None, size=10, font="Arial"):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font
            run.font.size = Pt(size)
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)


def add_para(doc: Document, text: str, *, bold=False, color=DARK, size=10, italic=False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)


def add_badge(paragraph, text: str, status: str) -> None:
    """Append a status badge. status ∈ {'ok', 'warn', 'danger'}"""
    if status == "ok":
        color, shade = GREEN, SHADE_SUCCESS
    elif status == "warn":
        color, shade = ORANGE, SHADE_WARNING
    else:
        color, shade = RED, SHADE_DANGER
    run = paragraph.add_run(f"  ● {text}  ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = color
    # Shading on the run via rPr
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), shade)
    rPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        set_cell_shading(cell, SHADE_NAVY)
        set_cell_font(cell, bold=True, color=WHITE, size=10)

    # Body rows with alternating shading
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            shade = SHADE_ALT_ROW if r_idx % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, shade)
            set_cell_font(cell, bold=False, color=DARK, size=9)

    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for i, width_cm in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width_cm)


def add_cover_page(doc: Document) -> None:
    # Top spacer
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BUILD PLAN")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = MEDIUM

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("\nSkillForge AI")
    run.font.name = "Arial"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("AI-Powered Employee Skill Assessment Platform")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("Execution Plan • Phase 1 Hyper-MVP to Phase 3 SaaS Launch")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MEDIUM

    # Metadata table
    doc.add_paragraph()
    doc.add_paragraph()
    meta = [
        ["Document ID", DOC_ID],
        ["Version", VERSION],
        ["Date", DOC_DATE],
        ["Prepared By", f"{AUTHOR} ({COMPANY})"],
        ["Prepared For", FOR_ORG],
        ["Classification", "Internal — Confidential"],
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta):
        lc, vc = table.rows[i].cells
        lc.text = label
        vc.text = value
        set_cell_shading(lc, SHADE_NAVY)
        set_cell_font(lc, bold=True, color=WHITE, size=10)
        set_cell_shading(vc, "FFFFFF")
        set_cell_font(vc, bold=False, color=DARK, size=10)
        lc.width = Cm(4.5)
        vc.width = Cm(10.5)

    # Classification notice
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("⚠ INTERNAL / CONFIDENTIAL — Not for external distribution")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_page_break()


def build_document() -> Document:
    doc = Document()

    # Page setup — US Letter, 1-inch margins (per global CLAUDE.md)
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    add_cover_page(doc)

    # ----- 1. Document Information -----
    add_heading(doc, "1. Document Information", level=1)
    add_table(
        doc,
        headers=["Field", "Value"],
        rows=[
            ["Document Name", "SkillForge AI — Build Plan"],
            ["Document ID", DOC_ID],
            ["Version", VERSION],
            ["Date", DOC_DATE],
            ["Owner", AUTHOR],
            ["Stakeholders", "Qualtech Product + Engineering leadership, HR Lead, Security Lead"],
            ["Classification", "Internal — Confidential"],
        ],
        col_widths_cm=[5, 11],
    )

    # ----- 2. Version History -----
    add_heading(doc, "2. Version History", level=1)
    add_table(
        doc,
        headers=["Version", "Date", "Author", "Changes"],
        rows=[
            ["v1.0", DOC_DATE, AUTHOR, "Initial execution plan — hyper-MVP + full Phase 1 + Phase 2/3 outlines"],
        ],
        col_widths_cm=[2, 2.5, 4, 7.5],
    )

    # ----- 3. Executive Summary -----
    add_heading(doc, "3. Executive Summary", level=1)
    add_para(
        doc,
        "SkillForge AI must support the Qualtech April–June 2026 appraisal cycle, with employee "
        "assessments due by end of May 2026. From today (2026-04-18), we have approximately six "
        "calendar weeks — equivalent to three two-week sprints — before the appraisal window closes.",
    )
    add_para(
        doc,
        "This plan defines a three-track build: a Hyper-MVP (Weeks 1–6) delivering only the "
        "assessment flow needed to run the cycle, a Full Phase 1 (Weeks 1–16) delivering all P0 and "
        "P1 features from the strategic plan §4.1, and a Phase 2+3 roadmap (Weeks 17–48) covering "
        "AI-powered evaluation and SaaS productization per plan §4.2 and §4.3.",
    )
    add_para(
        doc,
        "All 10 open architectural decisions (captured in OPEN_DECISIONS.md) must be resolved "
        "during the two-day Architecture Workshop on 2026-04-20 and 2026-04-21. Sprint 0 (Week 0) "
        "delivers a fully running dev environment by 2026-04-24 so that Sprint 1 execution begins "
        "Monday 2026-04-27.",
    )

    # ----- 4. Dashboard View -----
    add_heading(doc, "4. Dashboard View", level=1)
    dash = doc.add_paragraph()
    add_badge(dash, "Hyper-MVP On Track", "ok")
    add_badge(dash, "10 Decisions Open", "warn")
    add_badge(dash, "Team Not Yet Assembled", "danger")

    add_para(doc, "")
    add_table(
        doc,
        headers=["Metric", "Current", "Target (Sprint 3)", "Status"],
        rows=[
            ["Person-days committed (Hyper-MVP)", "95 / 270", "95 / 270 used", "● Healthy buffer"],
            ["P0 features locked", "6 / 6 scoped", "6 deployed", "● Scoped"],
            ["Open architectural decisions", "10", "0", "● Must close by 2026-04-21"],
            ["Team assembled", "Pending", "9–10 FTE", "● At risk"],
        ],
        col_widths_cm=[5.5, 3.5, 3.5, 3.5],
    )

    # ----- 5. Completed Work -----
    add_heading(doc, "5. Completed Work", level=1)
    add_para(doc, "Pre-sprint foundation (2026-04-18):", bold=True)
    add_table(
        doc,
        headers=["Artifact", "Description", "Version", "Date"],
        rows=[
            ["Strategic Plan DOCX", "Product vision, scope, stakeholders, success criteria", "v1.0", "2026-04-01"],
            [".claude/ setup", "8 project skills + 5 hooks + settings.json + CLAUDE.md", "v1.0", "2026-04-18"],
            ["Project memory files", "9 memory files covering overview, phases, tech stack, data model, security, feedback", "v1.0", "2026-04-18"],
            ["BUILD_PLAN.md", "This document — 3-track execution plan", "v1.0", "2026-04-18"],
            ["OPEN_DECISIONS.md", "10 decisions blocking Sprint 1 kickoff", "v1.0", "2026-04-18"],
            ["Sprint 0 Checklist", "Day-by-day checklist for Week 0 foundation", "v1.0", "2026-04-18"],
        ],
        col_widths_cm=[3.5, 8, 1.8, 2.5],
    )

    # ----- 6. Work in Progress -----
    add_heading(doc, "6. Work in Progress", level=1)
    add_table(
        doc,
        headers=["Feature", "Description", "Owner", "Expected Completion"],
        rows=[
            ["Architecture Workshop", "Resolve 10 open decisions, draft ADRs", "Tech Lead", "2026-04-21"],
            ["Repo scaffold", "Monorepo with apps/ services/ packages/ infra/ prompts/", "Tech Lead + DevOps", "2026-04-22"],
            ["CI/CD pipeline", "GitHub Actions with typecheck/lint/test/build/deploy", "DevOps", "2026-04-23"],
            ["Design system", "shadcn/ui tokens, Figma file, wireframes for Sprint 1 screens", "UI/UX + FE Lead", "2026-04-24"],
            ["DB schema v0 + auth skeleton", "Core entities, RBAC, JWT flow", "BE Lead", "2026-04-24"],
        ],
        col_widths_cm=[3.5, 7, 3, 2.5],
    )

    # ----- 7. Pending Items -----
    add_heading(doc, "7. Pending Items", level=1)
    add_table(
        doc,
        headers=["Item", "Description", "Priority", "Remarks"],
        rows=[
            ["Team assembly", "Allocate/hire 9-10 FTEs per plan §5.1", "P0", "Blocking Sprint 0 start"],
            ["Qualtech IdP discovery", "Confirm Azure AD / Okta / other; SAML/OIDC capability", "P0", "Needed for D9 Auth decision"],
            ["Appraisal-system export spec", "Column layout for CSV export Qualtech HR will consume", "P0", "Needed for Sprint 3"],
            ["Pilot user list", "20 employees + 5 managers for UAT", "P0", "Needed for Sprint 3"],
            ["Competency framework content", "Actual AI maturity levels + role mappings for Qualtech", "P0", "Needed for Sprint 1 seed"],
        ],
        col_widths_cm=[3.5, 7, 1.5, 4],
    )

    # ----- 8. Future Roadmap -----
    add_heading(doc, "8. Future Roadmap", level=1)
    add_table(
        doc,
        headers=["Phase", "Feature Bucket", "Timeline", "Effort (person-days)"],
        rows=[
            ["Phase 1 — Hyper-MVP", "Self-assessment + manager scoring + CSV export", "Weeks 1–6", "~95"],
            ["Phase 1 — Fast-follow", "Dashboards, SSO, notifications, reporting", "Weeks 7–16", "~72"],
            ["Phase 2", "AI analysis, peer review, HRMS integration, mobile, learning paths, bias detection", "Weeks 17–32", "165"],
            ["Phase 3", "Multi-tenant provisioning, billing, marketplace, SOC 2 prep, 10K-user scale", "Weeks 33–48", "150"],
        ],
        col_widths_cm=[4, 8, 2.5, 2.5],
    )

    # ----- 9. Upcoming Deliveries -----
    add_heading(doc, "9. Upcoming Deliveries", level=1)
    add_table(
        doc,
        headers=["Milestone", "Features", "Planned Release"],
        rows=[
            ["Sprint 0 done", "Dev environment operational", "2026-04-24 (Fri)"],
            ["Sprint 1 done", "Auth + users + framework + HR admin UI", "2026-05-01 (Fri)"],
            ["Sprint 2 done", "Self-assessment end-to-end, artifact upload, reminders", "2026-05-15 (Fri)"],
            ["Sprint 3 done", "Manager scoring, composite score, CSV export, pilot UAT", "2026-05-29 (Fri)"],
            ["Hyper-MVP Production", "Qualtech runs appraisal cycle on SkillForge", "2026-06-01 (Mon)"],
            ["Full Phase 1", "All P0/P1 Phase-1 features delivered", "2026-08-08 (Week 16)"],
            ["Phase 2 Production", "AI-assisted evaluation live", "2026-11-28 (Week 32)"],
            ["Phase 3 — First external client", "Qualtech + 1 external pilot", "2027-03-20 (Week 48)"],
        ],
        col_widths_cm=[4, 8.5, 4],
    )

    # ----- 10. Risks & Dependencies -----
    add_heading(doc, "10. Risks & Dependencies", level=1)
    add_table(
        doc,
        headers=["Risk", "Severity", "Mitigation"],
        rows=[
            ["6-week MVP window may not absorb all P0", "High", "Hyper-MVP strips to 6 features; dashboards deferred to Sprint 4–5"],
            ["Team not yet assembled", "High", "Escalate to engineering leadership immediately; interim Tech Lead by 2026-04-19 minimum"],
            ["AI hallucination risk in Phase 2", "High", "Manager override mandatory; confidence thresholds; fallback to manual"],
            ["HRMS integration complexity", "Medium", "Phase 1 uses CSV; API integration starts with 2-week discovery in Sprint 9"],
            ["Tenant data isolation leak", "High", "sf-tenant-check skill + pre-edit hook enforce org_id filters from Day 1"],
            ["PII leak via Claude API", "High", "pre-edit-prompt-pii hook blocks raw PII in prompt files"],
            ["Qualtech IdP unknown", "Medium", "Kickoff discovery meeting in Day 1 of Sprint 0"],
            ["Scope creep", "Medium", "P0/P1/P2 discipline; change requests become ADRs; no feature merges without ticket"],
        ],
        col_widths_cm=[6, 2, 8.5],
    )

    # ----- 11. Architecture Overview -----
    add_heading(doc, "11. Architecture Overview", level=1)
    add_para(doc, "Stack (from plan §3.2 and reference_tech_stack memory):", bold=True)
    add_table(
        doc,
        headers=["Layer", "Technology"],
        rows=[
            ["Frontend (Web)", "Next.js 14+ App Router, TypeScript, Tailwind, shadcn/ui"],
            ["Frontend (Mobile, Phase 2)", "React Native (Expo)"],
            ["API Gateway", "Kong / AWS API Gateway"],
            ["Backend (core services)", "NestJS (TypeScript)"],
            ["Backend (enterprise integrations)", "Spring Boot (Java)"],
            ["Database", "PostgreSQL 15+ (primary), Redis 7 (cache/sessions)"],
            ["AI/LLM", "Claude API + LangChain; prompt caching enabled"],
            ["File Storage", "AWS S3 (prod) / MinIO (local)"],
            ["Search", "Elasticsearch / OpenSearch (Phase 2)"],
            ["Auth", "Keycloak or Auth0 (decision D9) — SAML 2.0 + OIDC"],
            ["Infrastructure", "AWS (ECS/EKS), Terraform, GitHub Actions"],
            ["Monitoring", "Grafana + Prometheus + Sentry"],
        ],
        col_widths_cm=[6, 10],
    )

    # ----- 12. Security Overview -----
    add_heading(doc, "12. Security Overview", level=1)
    add_table(
        doc,
        headers=["Control", "Standard", "Status"],
        rows=[
            ["Tenant isolation", "SOC 2 Type II", "● Enforced via hook + skill"],
            ["PII stripped before Claude calls", "Responsible AI / DPDP Act 2023", "● Enforced via pre-edit hook"],
            ["Audit log (append-only)", "ISO 27001 A.12.4", "● Designed; build Sprint 1"],
            ["RBAC with tenant-scoped permissions", "Least Privilege", "● Designed"],
            ["JWT short-expiry + refresh rotation", "OWASP ASVS L2", "● Designed"],
            ["Encryption at rest (RDS + KMS)", "NIST SP 800-175B", "● Planned Sprint 0"],
            ["TLS 1.3 on all endpoints", "OWASP ASVS L2", "● Planned Sprint 0"],
            ["SAST + dependency scan in CI", "OWASP Testing Guide v4", "● Planned Sprint 0"],
        ],
        col_widths_cm=[7, 5.5, 3.5],
    )

    # ----- 13. VAPT/Vulnerability Findings -----
    add_heading(doc, "13. VAPT / Vulnerability Findings", level=1)
    add_para(doc, "Not yet run — internal scan scheduled Sprint 7, external pentest in Sprint 15 (Phase 2 close).", italic=True)

    # ----- 14. Functional Gaps Summary -----
    add_heading(doc, "14. Functional Gaps Summary", level=1)
    add_table(
        doc,
        headers=["Phase", "Effort (days)", "Focus"],
        rows=[
            ["Hyper-MVP (W1–6)", "~95", "6 features minimum for appraisal cycle"],
            ["Phase 1 remainder (W7–16)", "~72", "Dashboards, SSO, notifications, reporting"],
            ["Phase 2 (W17–32)", "165", "AI, peer review, HRMS, mobile, learning"],
            ["Phase 3 (W33–48)", "150", "Multi-tenancy, billing, marketplace, scale"],
        ],
        col_widths_cm=[5, 3, 8],
    )

    # ----- 15. Environment Access -----
    add_heading(doc, "15. Environment Access", level=1)
    add_table(
        doc,
        headers=["Environment", "URL", "Purpose"],
        rows=[
            ["Local", "http://localhost:3000", "Each developer's machine"],
            ["Dev", "https://dev.skillforge.qualtech.internal (TBD)", "Shared branch integration"],
            ["Staging", "https://staging.skillforge.qualtech.internal (TBD)", "Pre-prod, UAT"],
            ["Production", "https://skillforge.qualtech.internal (TBD)", "Qualtech workforce"],
        ],
        col_widths_cm=[3.5, 7, 5.5],
    )

    # ----- 16. Effort Summary -----
    add_heading(doc, "16. Effort Summary", level=1)
    add_table(
        doc,
        headers=["Metric", "Value"],
        rows=[
            ["Total team size (Phase 1)", "9–10 FTE"],
            ["Phase 1 duration", "16 weeks (8 sprints of 2 weeks)"],
            ["Phase 1 effort", "167 person-days"],
            ["Hyper-MVP window", "6 weeks (3 sprints)"],
            ["Hyper-MVP effort", "~95 person-days"],
            ["Total programme duration", "48 weeks (Phase 1 + 2 + 3)"],
            ["Total programme effort", "~482 person-days"],
        ],
        col_widths_cm=[6, 10],
    )

    # ----- 17. Commit Type Breakdown -----
    add_heading(doc, "17. Commit Type Breakdown (planned)", level=1)
    add_table(
        doc,
        headers=["Type", "Hyper-MVP est.", "Phase 1 total", "Phase 2 total"],
        rows=[
            ["feat", "~45", "~80", "~85"],
            ["fix", "~15", "~30", "~35"],
            ["docs", "~8", "~15", "~20"],
            ["security", "~5", "~10", "~20"],
            ["test", "~15", "~25", "~40"],
            ["chore / refactor", "~12", "~20", "~25"],
        ],
        col_widths_cm=[4, 4, 4, 4],
    )

    # ----- 18. References & Links -----
    add_heading(doc, "18. References & Links", level=1)
    add_table(
        doc,
        headers=["Reference", "Location"],
        rows=[
            ["Strategic plan (source of truth)", "SkillForge_AI_Project_Plan.docx"],
            ["Build plan (this doc's Markdown)", "BUILD_PLAN.md"],
            ["Sprint 0 checklist", "docs/SPRINT_0_CHECKLIST.md"],
            ["Open decisions", "OPEN_DECISIONS.md"],
            ["Project memory index", "~/.claude/projects/-Users-aksatyam-SelfWork-SkillForge/memory/MEMORY.md"],
            ["Project Claude guide", "CLAUDE.md"],
            ["Project-scoped skills", ".claude/skills/"],
            ["Project hooks", ".claude/hooks/"],
            ["Repository (after Sprint 0)", "git@github.com:qualtech/skillforge.git (TBD)"],
            ["Jira / Linear board", "TBD in Sprint 0"],
        ],
        col_widths_cm=[5, 11],
    )

    # Footer block
    doc.add_page_break()
    add_heading(doc, "Prepared By", level=2)
    add_para(doc, f"{AUTHOR}", bold=True, color=NAVY, size=12)
    add_para(doc, COMPANY, color=DARK, size=10)
    add_para(doc, f"For: {FOR_ORG}", color=DARK, size=10)
    add_para(doc, f"Date: {DOC_DATE}", color=MEDIUM, size=9)
    add_para(doc, "Classification: Internal — Confidential", color=RED, size=9, bold=True)

    return doc


def main() -> None:
    out_path = Path(__file__).parent.parent / "SkillForge_AI_Build_Plan.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Wrote {out_path} ({out_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
