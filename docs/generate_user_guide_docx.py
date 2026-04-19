#!/usr/bin/env python3
"""
SkillForge AI — User Guide DOCX generator.

Generates a branded enterprise-style User Guide from the content of
docs/USER_GUIDE.md. Follows the global brand standard defined in
~/.claude/CLAUDE.md:

  - Brand colors:  Navy #1B3A5C, Blue #2E75B6, Green #27AE60, Orange #E67E22,
                   Red #E74C3C, Dark #2C3E50, Medium #7F8C8D, Light #F0F4F8
  - Fonts:         Arial body, Consolas code
  - Page:          US Letter, 1-inch margins
  - Tables:        WidthType.DXA, ShadingType.CLEAR, alternating #F8F9FA/#FFFFFF
  - Header:        right-aligned doc title + version
  - Footer:        centered page numbers + navy prepared-by block
  - Classification: "Internal / Confidential" in red on cover

Attribution:
  Author   — Ashish Kumar Satyam
  For      — Qualtech

Run:
  python3 docs/generate_user_guide_docx.py

Output:
  docs/SkillForge-User-Guide-v1.0.docx
"""

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ── Brand palette ─────────────────────────────────────────────────────

NAVY     = RGBColor(0x1B, 0x3A, 0x5C)
BLUE     = RGBColor(0x2E, 0x75, 0xB6)
GREEN    = RGBColor(0x27, 0xAE, 0x60)
ORANGE   = RGBColor(0xE6, 0x7E, 0x22)
RED      = RGBColor(0xE7, 0x4C, 0x3C)
DARK     = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM   = RGBColor(0x7F, 0x8C, 0x8D)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "F0F4F8"
ALT_ROW  = "F8F9FA"
SUCCESS_BG = "E8F5E9"
WARN_BG    = "FFF8E1"
DANGER_BG  = "FFEBEE"
HEADER_BG  = "1B3A5C"   # navy for table header rows

# ── Doc metadata ──────────────────────────────────────────────────────

DOC_TITLE     = "SkillForge AI — Application User Guide"
DOC_VERSION   = "1.0"
DOC_DATE      = date(2026, 4, 19)
DOC_OWNER     = "Qualtech"
DOC_AUTHOR    = "Ashish Kumar Satyam"
DOC_FOR       = "Qualtech"
DOC_CLASS     = "Internal / Confidential"

OUT_PATH = Path(__file__).parent / "SkillForge-User-Guide-v1.0.docx"


# ═══════════════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════════════

def shade(cell, hex_color: str) -> None:
    """Apply a solid background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run(run, *, font="Arial", size=10, color=DARK, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Ensure east-asian and cs fonts don't override
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)


def add_para(
    doc_or_cell,
    text: str = "",
    *,
    font="Arial",
    size=10,
    color=DARK,
    bold=False,
    italic=False,
    align=None,
    space_before=0,
    space_after=4,
    line_spacing=1.15,
):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        set_run(r, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int = 1):
    """Levels: 1 = chapter (H1), 2 = section (H2), 3 = subsection (H3)."""
    sizes = {1: 22, 2: 16, 3: 13}
    spacing_before = {1: 18, 2: 12, 3: 8}
    spacing_after = {1: 10, 2: 6, 3: 4}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before.get(level, 6))
    p.paragraph_format.space_after = Pt(spacing_after.get(level, 4))
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 12), color=NAVY, bold=True)
    if level == 1:
        # thin rule under chapter heading
        p_border = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1B3A5C")
        pbdr.append(bottom)
        p_border.append(pbdr)
    return p


def add_code(doc_or_cell, text: str):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    for line in text.rstrip().split("\n"):
        if line:
            r = p.add_run(line)
            set_run(r, font="Consolas", size=9, color=DARK)
        p.add_run().add_break(WD_BREAK.LINE)
    # shade the paragraph? python-docx doesn't have paragraph-shading API; use a
    # one-cell 1x1 table instead if backgrounding becomes important later.


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.4 * level)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run(r, size=10, color=DARK)
    return p


def add_callout(doc, label: str, text: str, *, color=BLUE, bg=LIGHT_BG):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, bg)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(f"{label}  ")
    set_run(r1, size=10, color=color, bold=True)
    r2 = p1.add_run(text)
    set_run(r2, size=10, color=DARK)
    add_para(doc, "", space_after=6)


def add_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_inches: list[float] | None = None,
):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.allow_autofit = False

    if col_widths_inches is None:
        total = 6.5
        col_widths_inches = [total / len(headers)] * len(headers)

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.width = Inches(col_widths_inches[i])
        shade(cell, HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        r = cell.paragraphs[0].add_run(h)
        set_run(r, size=10, color=WHITE, bold=True)

    # Body rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx + 1, c_idx)
            cell.width = Inches(col_widths_inches[c_idx])
            if r_idx % 2 == 1:
                shade(cell, ALT_ROW)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            # Support a status-pill convention: "● Completed" → color the dot
            val_str = str(val)
            if val_str.startswith("●"):
                parts = val_str.split(" ", 1)
                dot_color = GREEN
                if "At Risk" in val_str or "In Progress" in val_str or "Warning" in val_str:
                    dot_color = ORANGE
                elif "Delayed" in val_str or "Blocked" in val_str or "Pending" in val_str or "Danger" in val_str:
                    dot_color = RED
                elif "Completed" in val_str or "Done" in val_str or "On Track" in val_str or "Yes" in val_str:
                    dot_color = GREEN
                r_dot = cell.paragraphs[0].add_run(parts[0] + " ")
                set_run(r_dot, size=10, color=dot_color, bold=True)
                if len(parts) > 1:
                    r_text = cell.paragraphs[0].add_run(parts[1])
                    set_run(r_text, size=10, color=DARK)
            else:
                r_run = cell.paragraphs[0].add_run(val_str)
                set_run(r_run, size=10, color=DARK)
    add_para(doc, "", space_after=6)
    return tbl


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)


def set_header_footer(section):
    # Right-aligned header: doc title + version
    hdr = section.header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hdr_run = hdr.add_run(f"{DOC_TITLE}  ·  v{DOC_VERSION}")
    set_run(hdr_run, size=9, color=MEDIUM)

    # Centered footer with page numbers
    ftr = section.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ftr.add_run()
    set_run(run, size=9, color=MEDIUM)
    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

    run2 = ftr.add_run(" / ")
    set_run(run2, size=9, color=MEDIUM)

    run3 = ftr.add_run()
    set_run(run3, size=9, color=MEDIUM)
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld_sep2 = OxmlElement("w:fldChar")
    fld_sep2.set(qn("w:fldCharType"), "separate")
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_begin2)
    run3._r.append(instr2)
    run3._r.append(fld_sep2)
    run3._r.append(fld_end2)


# ═══════════════════════════════════════════════════════════════════════
#  Document body — Cover + Metadata
# ═══════════════════════════════════════════════════════════════════════

def write_cover(doc):
    # Big navy top block for cover
    add_para(doc, "", space_before=60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SkillForge AI")
    set_run(r, size=32, color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Application User Guide")
    set_run(r, size=22, color=BLUE, bold=False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    r = p3.add_run("For all roles — Employees, Managers, HR Admins, AI Champions, Leadership, Super Admins")
    set_run(r, size=11, color=MEDIUM, italic=True)

    add_para(doc, "", space_before=40)
    # Classification banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(4)
    shade(cell, "FFEBEE")
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f"  {DOC_CLASS}  ")
    set_run(r, size=11, color=RED, bold=True)

    add_para(doc, "", space_before=30)

    # Metadata table
    rows = [
        ["Document name", DOC_TITLE],
        ["Version", DOC_VERSION],
        ["Date", DOC_DATE.isoformat()],
        ["Owner", DOC_OWNER],
        ["Author", DOC_AUTHOR],
        ["Prepared for", DOC_FOR],
        ["Classification", DOC_CLASS],
        ["Phase", "Phase 1 (Hyper-MVP) — through Sprint 6"],
        ["Distribution", "Internal — SkillForge platform users + admins"],
    ]
    meta = doc.add_table(rows=len(rows), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    for i, (k, v) in enumerate(rows):
        c1 = meta.cell(i, 0)
        c2 = meta.cell(i, 1)
        c1.width = Inches(1.8)
        c2.width = Inches(4.0)
        if i % 2 == 1:
            shade(c1, ALT_ROW)
            shade(c2, ALT_ROW)
        c1.paragraphs[0].paragraph_format.space_before = Pt(3)
        c1.paragraphs[0].paragraph_format.space_after = Pt(3)
        c2.paragraphs[0].paragraph_format.space_before = Pt(3)
        c2.paragraphs[0].paragraph_format.space_after = Pt(3)
        r1 = c1.paragraphs[0].add_run(k)
        set_run(r1, size=10, color=NAVY, bold=True)
        r2 = c2.paragraphs[0].add_run(v)
        set_run(r2, size=10, color=DARK)

    add_para(doc, "", space_before=40)

    # "Prepared By" footer block on cover
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, HEADER_BG)
    for line, bold in [
        ("Prepared by", False),
        (DOC_AUTHOR, True),
        (f"for {DOC_FOR}", False),
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run(r, size=11, color=WHITE, bold=bold)
    # Remove default empty paragraph from the cell
    first = cell.paragraphs[0]
    if first.text == "":
        first._element.getparent().remove(first._element)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
#  Document body — Version history
# ═══════════════════════════════════════════════════════════════════════

def write_version_history(doc):
    add_heading(doc, "Version history", level=1)
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "2026-04-19", DOC_AUTHOR, "Initial release — Phase 1, through Sprint 6. Covers all 6 roles."],
        ],
        col_widths_inches=[0.8, 1.1, 1.8, 2.8],
    )
    add_para(doc, "", space_after=10)


# ═══════════════════════════════════════════════════════════════════════
#  Document body — Table of contents (static summary)
# ═══════════════════════════════════════════════════════════════════════

def write_toc(doc):
    add_heading(doc, "Table of contents", level=1)
    chapters = [
        ("1", "Welcome to SkillForge AI", "What the platform does + five platform promises"),
        ("2", "Getting started", "Invite-accept, sign-in, SSO, MFA, nav, profile"),
        ("3", "For Employees", "Self-assessments, artifacts, scorecard"),
        ("4", "For Managers", "Roster, team overview, scoring form, overrides"),
        ("5", "For HR Admins", "Cycles, users, frameworks, reports, exports"),
        ("6", "For AI Champions", "Framework quality + Phase-2 calibration"),
        ("7", "For Leadership", "Read-only org capability views"),
        ("8", "For Super Admins", "Break-glass + tenant ops + audit"),
        ("9", "Notifications", "Settings + what you can/can't opt out of"),
        ("10", "End-to-End Journey", "Single-cycle flow across all roles"),
        ("11", "Security, Privacy, and Your Data", "Auth, multi-tenancy, encryption, AI+PII, DPDP"),
        ("12", "What's coming in Phase 2", "AI intelligence + mobile + HRMS"),
        ("13", "Troubleshooting", "Common failures + resolutions"),
        ("14", "FAQ", "Top 10 questions"),
        ("15", "Glossary", "28 terms you'll see in the product"),
        ("16", "Support and contact", "Who to ping for what"),
    ]
    add_table(
        doc,
        ["§", "Chapter", "What it covers"],
        [list(c) for c in chapters],
        col_widths_inches=[0.4, 2.2, 3.9],
    )
    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
#  Document body — Chapter writers
# ═══════════════════════════════════════════════════════════════════════

def ch_1_welcome(doc):
    add_heading(doc, "1. Welcome to SkillForge AI", level=1)
    add_para(
        doc,
        "SkillForge AI is Qualtech's internal platform for capturing, scoring, and analyzing "
        "employee AI-capability across a structured competency framework. It replaces scattered "
        "spreadsheets and one-off review forms with a single, auditable workflow that feeds "
        "appraisals, training plans, and organizational capability reports.",
    )

    add_heading(doc, "1.1 What SkillForge gives you", level=2)
    add_table(
        doc,
        ["If you are a…", "You use SkillForge to…"],
        [
            ["Employee", "Self-assess against your role's framework, upload supporting artifacts, view your scorecard, track progress across cycles."],
            ["Manager", "Review reports' self-assessments, score each dimension with rationale, override AI suggestions (Phase 2) where required."],
            ["HR Admin", "Publish frameworks, open / lock / finalize / close review cycles, invite users, pull reports and CSV exports."],
            ["AI Champion", "Curate frameworks, monitor score distributions, own the dimension-level rubrics."],
            ["Leadership", "Read dashboards and org-wide capability reports without changing underlying data."],
            ["Super Admin", "Operate the tenant — everything above, plus role changes, SSO config, break-glass access."],
        ],
        col_widths_inches=[1.3, 4.8],
    )

    add_heading(doc, "1.2 The five promises SkillForge keeps", level=2)
    for i, text in enumerate(
        [
            ("One source of truth.", " Every score — self, manager, AI-suggested (Phase 2), composite — is stored once, timestamped, and audited."),
            ("Manager decisions rule.", " AI scores are advisory only. When they exist (Phase 2), a manager must confirm or override with a written rationale before the score is final."),
            ("Your data stays in your tenant.", " Qualtech's data never mixes with any other organization on the platform. Queries filter by organization at the database level."),
            ("Audit-first.", " Who scored what, when, and what they overrode — all captured, all retainable for seven years."),
            ("Secure by default.", " OWASP ASVS L2, SOC 2 Type II alignment, DPDP Act 2023 compliance. PII stripped before any AI call."),
        ],
        start=1,
    ):
        bold_part, rest = text
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{i}. ")
        set_run(r1, size=10, color=DARK, bold=True)
        r2 = p.add_run(bold_part)
        set_run(r2, size=10, color=NAVY, bold=True)
        r3 = p.add_run(rest)
        set_run(r3, size=10, color=DARK)


def ch_2_getting_started(doc):
    add_heading(doc, "2. Getting started", level=1)

    add_heading(doc, "2.1 Accepting your invite", level=2)
    add_para(doc, "Your HR Admin creates your account and emails you a one-time invite link:")
    add_code(doc, "https://skillforge.qualtech.ai/invite/<long-random-token>")
    add_para(doc, "When you click, the invite-accept page:")
    add_bullet(doc, "Pre-fills your name and email (raise a ticket if wrong — not editable here).")
    add_bullet(doc, "Shows a Set-password form. Min 8 chars, letter + number. Passphrases preferred.")
    add_bullet(doc, "Activates your account, signs you in, and takes you to /dashboard.")
    add_callout(
        doc,
        "Why this flow is atomic:",
        "the server treats 'consume invite → set password → issue first session' as one transaction. "
        "Either all three succeed or none do — you can't end up with a half-activated account.",
        color=BLUE,
        bg=LIGHT_BG,
    )
    add_callout(
        doc,
        "Heads up:",
        "Invite links expire 7 days after issue. Expired = ask HR to re-invite (old token is burned). "
        "Single-use — a second visit says 'already used'.",
        color=ORANGE,
        bg=WARN_BG,
    )

    add_heading(doc, "2.2 Signing in", level=2)
    add_heading(doc, "Email + password", level=3)
    add_bullet(doc, "Go to /login. Enter email + password. Submit.")
    add_bullet(doc, "On success: lands on /dashboard with HTTP-only, SameSite=Lax cookies.")
    add_bullet(doc, "Rate-limited after ~10 failed attempts/minute/IP. Wait and retry.")
    add_heading(doc, "Single Sign-On (SSO)", level=3)
    add_bullet(doc, "Click 'Sign in with SSO' → your IdP (Keycloak / Azure AD / Okta).")
    add_bullet(doc, "Complete your usual corporate login (including MFA).")
    add_bullet(doc, "Bounce back to SkillForge, already signed in.")
    add_callout(
        doc,
        "Why SSO is safer:",
        "SkillForge never sees your corporate password. Your IdP sends a signed assertion of "
        "your identity pinned to our app's audience claim.",
        color=GREEN,
        bg=SUCCESS_BG,
    )
    add_heading(doc, "Multi-Factor Authentication (MFA)", level=3)
    add_para(doc, "Mandatory for HR Admin, AI Champion, Super Admin. Optional for others.")
    add_bullet(doc, "Profile menu → Security → Enable MFA")
    add_bullet(doc, "Scan QR with Google Authenticator / Authy / 1Password")
    add_bullet(doc, "Enter 6-digit code")
    add_bullet(doc, "Save recovery codes (used if you lose your phone)")

    add_heading(doc, "2.3 Roles at a glance", level=2)
    add_para(
        doc,
        "SkillForge has six roles, assigned by an HR Admin at invite time. Your current role is "
        "displayed in the sidebar under your name.",
    )
    add_table(
        doc,
        ["Role", "Primary use", "Can create / edit…", "Read-only access to…"],
        [
            ["employee", "Self-assess, view own scorecard", "Own assessments (until submit)", "Own history"],
            ["manager", "Score direct reports", "Manager scoring for own team", "Own team roster + history"],
            ["hr_admin", "Run cycles + frameworks", "Frameworks, cycles, users, templates", "Whole tenant"],
            ["ai_champion", "Curate frameworks", "Frameworks (same as HR Admin)", "Whole tenant"],
            ["leadership", "Read dashboards", "Nothing", "Whole tenant — read-only"],
            ["super_admin", "Tenant ops + break-glass", "Everything", "Everything"],
        ],
        col_widths_inches=[1.1, 1.6, 1.8, 1.6],
    )

    add_heading(doc, "Role matrix — who sees which nav item", level=3)
    yes = "●"
    no = "—"
    add_table(
        doc,
        ["Nav item", "emp", "mgr", "hr", "ai_c", "lead", "sup"],
        [
            ["Dashboard",             yes, yes, yes, yes, yes, yes],
            ["My Assessments",        yes, yes, yes, yes, yes, yes],
            ["My scorecard",          yes, yes, yes, yes, yes, yes],
            ["Team overview",         no,  yes, yes, no,  no,  yes],
            ["Team roster",           no,  yes, yes, no,  no,  yes],
            ["Frameworks",            no,  no,  yes, yes, no,  yes],
            ["Users",                 no,  no,  yes, no,  no,  yes],
            ["Cycles",                no,  no,  yes, no,  no,  yes],
            ["Reports",               no,  no,  yes, yes, no,  yes],
            ["Export templates",      no,  no,  yes, no,  no,  yes],
            ["HR Dashboard",          no,  no,  yes, no,  no,  yes],
            ["Notification settings", yes, yes, yes, yes, yes, yes],
        ],
        col_widths_inches=[2.0, 0.7, 0.7, 0.7, 0.8, 0.7, 0.7],
    )

    add_heading(doc, "2.4 Navigating the app shell", level=2)
    add_bullet(doc, "Sidebar left, main content right — same on every page.")
    add_bullet(doc, "Active nav item in navy; sub-pages keep the parent highlighted.")
    add_bullet(doc, "Organization name at top confirms the tenant.")
    add_bullet(doc, "Role pill at the bottom under your name = your source of truth for permissions.")
    add_bullet(doc, "Sign out revokes the refresh token server-side — it does not just hide the UI.")


def ch_3_employees(doc):
    add_heading(doc, "3. For Employees", level=1)
    add_callout(
        doc,
        "You are an employee if",
        "your role pill says EMPLOYEE. You own your self-assessments, view your scorecard, and "
        "upload artifacts that support your scores.",
    )

    add_heading(doc, "3.1 Your dashboard (/dashboard)", level=2)
    add_bullet(doc, "Open cycles — card per cycle accepting self-assessments, with Start/Resume CTA.")
    add_bullet(doc, "Upcoming deadline — soonest selfAssessmentDeadline across your open cycles.")
    add_bullet(doc, "Quick stats strip — completion %, last submitted date, last-cycle avg.")

    add_heading(doc, "3.2 My Assessments (/assessments)", level=2)
    add_para(doc, "Grid of your assessments, one card per cycle.")
    add_heading(doc, "Status grid", level=3)
    add_table(
        doc,
        ["Status", "What it means", "What you can do"],
        [
            ["not_started",       "Cycle opened; you haven't started",      "Click Start assessment"],
            ["self_submitted",    "You submitted; waiting on manager",      "View (read-only)"],
            ["manager_in_progress","Manager is actively scoring",           "View (read-only)"],
            ["ai_analyzed",       "AI has suggested scores (Phase 2)",      "View (read-only)"],
            ["manager_scored",    "Manager finalized their scores",         "View self + manager"],
            ["composite_computed","System computed weighted composite",     "View composite"],
            ["finalized",         "HR closed the cycle",                    "Read-only history"],
        ],
        col_widths_inches=[1.7, 2.6, 2.2],
    )

    add_heading(doc, "3.3 Filling out a self-assessment (/assessments/[id])", level=2)
    add_bullet(doc, "Left nav: dimension list. Right main: one card per dimension (rubric + score + comment + artifacts).")
    add_bullet(doc, "Score: 0.00 to 5.00, step 0.01. Fractionals allowed (e.g. 3.5).")
    add_bullet(doc, "Every dimension must be scored before Submit. Drafts can be partial.")
    add_bullet(doc, "Auto-save every 30 seconds, plus on every Save draft click.")
    add_bullet(doc, "Once submitted, form is locked. View-only.")
    add_bullet(doc, "Max 20 dimensions per assessment. Max 5 artifacts per dimension (25 MB each).")
    add_callout(
        doc,
        "Heads up:",
        "Uploading 20 large artifacts slows submission. Keep files under 10 MB where possible.",
        color=ORANGE,
        bg=WARN_BG,
    )

    add_heading(doc, "3.4 My scorecard (/scorecard)", level=2)
    add_bullet(doc, "Score radar chart — one polygon per cycle, overlaid, growth over time.")
    add_bullet(doc, "Score histogram — distribution of dimensions in latest finalized cycle.")
    add_bullet(doc, "Cycle-by-cycle table — self avg, manager avg, composite, delta.")
    add_bullet(doc, "Trajectory line — composite over time.")
    add_bullet(doc, "Framework filter if you've been assessed under multiple frameworks.")

    add_heading(doc, "3.5 Uploading artifacts", level=2)
    add_para(doc, "Every dimension card has an artifact slot. Evidence = stronger score defense.")
    add_bullet(doc, "PDF, PNG, JPG, MP4, DOCX, XLSX, TXT. Max 25 MB/file, max 5 per dimension.")
    add_bullet(doc, "Browser PUTs the file directly to S3 via a 15-minute signed URL carrying orgId.")
    add_bullet(doc, "Downloads use a 5-minute signed URL scoped to your organization.")
    add_bullet(doc, "Failed upload? Just re-drop the file. Fresh URL. Old one is single-use.")


def ch_4_managers(doc):
    add_heading(doc, "4. For Managers", level=1)
    add_callout(
        doc,
        "The one rule you need to remember:",
        "AI scores are advisory. Your number wins. Every time. This is a platform invariant and "
        "cannot be configured away.",
        color=NAVY,
        bg=LIGHT_BG,
    )

    add_heading(doc, "4.1 Your dashboard", level=2)
    add_bullet(doc, "Team pending count — reports who've submitted but you haven't scored.")
    add_bullet(doc, "Team completion donut — % of your team who've submitted.")
    add_bullet(doc, "Deadline strip — your own managerScoringDeadline for the current cycle.")

    add_heading(doc, "4.2 Team roster (/team)", level=2)
    add_bullet(doc, "Direct reports with current-cycle assessment status.")
    add_bullet(doc, "Pending (default) vs All filter.")
    add_bullet(doc, "Click row → manager scoring form for that person.")

    add_heading(doc, "4.3 Team overview (/team/overview)", level=2)
    add_bullet(doc, "Team radar — avg self vs avg manager, dual overlay reveals bias.")
    add_bullet(doc, "Per-person mini-cards — click for full scorecard.")
    add_bullet(doc, "Completion donut filterable by cycle.")

    add_heading(doc, "4.4 Scoring a report", level=2)
    add_para(doc, "Path: /team/[userId]/assessment/[assessmentId]")
    add_bullet(doc, "Header: employee, role, cycle, status, deadline countdown.")
    add_bullet(doc, "Per dimension: employee self-score (muted), their comment + artifacts, AI badge (Phase 2), manager input, rationale (20+ chars required).")
    add_bullet(doc, "Sticky sidebar: live composite preview using framework weights.")
    add_bullet(doc, "Overriding AI by >0.5 requires longer rationale and is audit-flagged.")
    add_bullet(doc, "Every dimension must be scored before Submit enables.")
    add_bullet(doc, "After Submit: status → manager_scored; not editable unless HR reopens.")

    add_heading(doc, "4.5 Common manager tasks", level=2)
    add_table(
        doc,
        ["Need to…", "How"],
        [
            ["Fix a submitted score",        "Ask HR Admin to reopen (audit-logged with your user id)"],
            ["Handle mid-scoring update",    "Banner appears; refresh — your draft is preserved"],
            ["Hand off a team",              "HR Admin re-parents reports; drafts travel with the assessment"],
            ["Score someone else's report",  "Cannot — server enforces direct-report relationship"],
        ],
        col_widths_inches=[2.2, 4.3],
    )


def ch_5_hr_admins(doc):
    add_heading(doc, "5. For HR Admins", level=1)
    add_para(doc, "You own the cycle. You also have org-wide read access and invite users.")

    add_heading(doc, "5.1 HR Dashboard (/hr)", level=2)
    add_bullet(doc, "KPI strip: open cycles, total users, avg completion %, avg scoring turnaround.")
    add_bullet(doc, "Live cycle cards with donuts + quick actions (Open / Lock / Bulk finalize / Close).")
    add_bullet(doc, "Recently closed (last 6). At-Risk pill on stalled cycles (48h+).")

    add_heading(doc, "5.2 Managing cycles", level=2)
    add_table(
        doc,
        ["Status", "What's allowed", "Transition to"],
        [
            ["draft",   "Add/remove participants. No assessments issued yet.", "open"],
            ["open",    "Employees self-assess; managers score; drafts editable.", "locked or closed"],
            ["locked",  "No new edits. HR can bulk-finalize composites.", "closed"],
            ["closed",  "Fully immutable. Read-only forever.",            "(terminal)"],
        ],
        col_widths_inches=[1.2, 3.7, 1.6],
    )
    add_para(doc, "Typical sequence:", bold=True, space_before=6)
    add_bullet(doc, "Create cycle (draft) → open → monitor → lock after deadline → bulk-finalize → close.")
    add_bullet(doc, "Bulk finalize ~2s per 100 assessments; flags unscored with Skip prompt.")
    add_bullet(doc, "Close is one-way. Reopen only at the per-assessment level after close.")

    add_heading(doc, "5.3 Managing users (/users)", level=2)
    add_bullet(doc, "Invite: email + role + manager + target maturity → 7-day one-time link.")
    add_bullet(doc, "Edit: role (can't demote self), manager, target. Email is immutable.")
    add_bullet(doc, "Deactivate: soft-delete. History retained; cannot sign in.")
    add_bullet(doc, "Resend invite rotates the token.")

    add_heading(doc, "5.4 Managing frameworks (/frameworks)", level=2)
    add_table(
        doc,
        ["Status", "Editable?", "Usable in cycles?"],
        [
            ["draft",     "Yes",                     "No"],
            ["published", "No (immutable snapshot)", "Yes"],
            ["archived",  "No",                      "No (old cycles keep snapshots)"],
        ],
        col_widths_inches=[1.3, 2.1, 3.1],
    )
    add_bullet(doc, "Publish is one-way. Clone-edit-publish to change anything.")
    add_bullet(doc, "Each cycle snapshots the framework — later edits don't retroactively change scores.")

    add_heading(doc, "5.5 CSV export + templates", level=2)
    add_bullet(doc, "RFC 4180 + UTF-8 BOM + fixed column order (matches appraisal import template).")
    add_bullet(doc, "/hr/templates lets you create custom column subsets + orders from an allowlist.")
    add_bullet(doc, "Disallowed fields (user.password, orgId, responsesJson) don't appear in picker — default-deny (see ADR-012).")
    add_bullet(doc, "Built-in templates are read-only. Tenant-custom templates are editable, but any 'builtin: true' flag is silently rewritten to false on save.")

    add_heading(doc, "5.6 Reports (/hr/reports)", level=2)
    add_bullet(doc, "Capability heat-map — dimension × role, colored by avg composite.")
    add_bullet(doc, "Cycle-over-cycle delta — per-dimension movement.")
    add_bullet(doc, "Role-target gap — % at or below target per dimension.")
    add_bullet(doc, "Outlier list — employees >2σ below role-peer avg.")


def ch_6_ai_champions(doc):
    add_heading(doc, "6. For AI Champions", level=1)
    add_para(doc, "Framework co-owner in Phase 1. Calibration + bias-detection owner in Phase 2.")

    add_heading(doc, "6.1 Phase-1 responsibilities", level=2)
    add_bullet(doc, "Rubric quality — each level's descriptor must be distinct and concrete.")
    add_bullet(doc, "Score-distribution audit — watch for statistical flatness, ceiling clustering.")
    add_bullet(doc, "Training-need feed — dimensions with ≥40% below target feed L&D.")

    add_heading(doc, "6.2 Screen access in Phase 1", level=2)
    add_table(
        doc,
        ["Screen", "Access"],
        [
            ["/frameworks + /frameworks/new + /frameworks/[id]", "Full edit"],
            ["/hr/reports",                                       "Read"],
            ["/users",                                            "Read-only in Phase 1"],
            ["/dashboard + /assessments + /scorecard",            "Full (also an employee)"],
            ["/cycles + /hr + /hr/cycles/[id]",                   "No access in Phase 1"],
            ["/hr/templates",                                     "No access in Phase 1"],
        ],
        col_widths_inches=[3.7, 2.5],
    )

    add_heading(doc, "6.3 Phase-2 preview", level=2)
    add_bullet(doc, "Calibration dashboard — manager-vs-AI delta heat-map, confidence-vs-accuracy curve.")
    add_bullet(doc, "Bias-detection panel — manager × role pairs with systematic skew.")
    add_bullet(doc, "Prompt monitoring — anonymized Claude input/output samples.")
    add_bullet(doc, "Rubric A/B tester — measure which descriptor variant reduces inter-rater variance.")


def ch_7_leadership(doc):
    add_heading(doc, "7. For Leadership", level=1)
    add_para(doc, "You read, you don't write. SkillForge is your organizational capability telescope.")

    add_heading(doc, "7.1 Access", level=2)
    add_table(
        doc,
        ["Screen", "What it's for"],
        [
            ["/dashboard", "Org-wide KPI strip — users, cycles, composite avg"],
            ["/scorecard + /assessments", "Your own — you're also an employee"],
            ["/team + /team/overview", "If you have direct reports, you get the manager view"],
            ["/settings/notifications", "Tune emails"],
        ],
        col_widths_inches=[2.5, 3.8],
    )

    add_heading(doc, "7.2 What you do NOT have", level=2)
    add_bullet(doc, "Cycle management (HR owns).")
    add_bullet(doc, "User invites or role changes (HR owns).")
    add_bullet(doc, "Framework editing (HR + AI Champion own).")
    add_bullet(doc, "CSV export (ask HR — every export is an audited event with a specific actor).")
    add_bullet(doc, "Override of manager scores (escalate through HR).")

    add_heading(doc, "7.3 Phase-2 Leadership dashboard", level=2)
    add_bullet(doc, "Org-wide capability heat-map over time (animated across cycles).")
    add_bullet(doc, "Role-group scorecards + roll-ups.")
    add_bullet(doc, "Training ROI tracker — cohort composite delta × headcount × salary.")
    add_bullet(doc, "Capability-to-strategy mapping.")


def ch_8_super_admins(doc):
    add_heading(doc, "8. For Super Admins", level=1)
    add_callout(
        doc,
        "Two rules:",
        "(1) Use your real role for day-to-day. Super Admin is break-glass. "
        "(2) Leave a trail — every action is audit-logged. SOC 2 expects monthly review.",
        color=RED,
        bg=DANGER_BG,
    )

    add_heading(doc, "8.1 What Super Admin adds on top of HR Admin", level=2)
    add_table(
        doc,
        ["Capability", "Where", "Used when"],
        [
            ["Change any user's role", "/users → edit", "New super admin; demoting after handoff"],
            ["Force-unlock a cycle", "/hr/cycles/[id] → Force unlock", "Post-lock issue; one more score needed"],
            ["Force-close a cycle",  "/hr/cycles/[id] → Force close",  "Emergency abort; data integrity"],
            ["Reopen any assessment","/hr/cycles/[id] → row → Reopen", "Regulatory requests, disputes"],
            ["Rotate JWT_SECRET",    "Infra (env variable + restart)", "Suspected token leak"],
            ["Rotate SSO_BRIDGE_SECRET","Infra (prod ≥32 chars)",      "Quarterly or on suspicion"],
            ["Audit log",            "DB direct (Phase 2: in-app viewer)", "Monthly SOC 2 review; incidents"],
        ],
        col_widths_inches=[2.0, 2.3, 2.2],
    )

    add_heading(doc, "8.2 Hard invariants you must NOT break", level=2)
    add_bullet(doc, "Tenant isolation — every DB query filters by org_id; RLS enforces at the engine level.")
    add_bullet(doc, "Audit append-only — no UPDATE/DELETE on audit_log for the service role.")
    add_bullet(doc, "AI advisory only — manager-score-over-AI-score is platform-level; no 'trust AI' toggle.")
    add_bullet(doc, "PII stripped before AI — the anonymizer runs server-side; no override.")
    add_bullet(doc, "No shared Super Admin credentials — every admin is one identifiable human.")


def ch_9_notifications(doc):
    add_heading(doc, "9. Notifications", level=1)
    add_heading(doc, "9.1 Notification types", level=2)
    add_table(
        doc,
        ["Notification", "Who gets it", "When"],
        [
            ["Cycle opened",              "Employees in the cycle", "HR Admin clicks Open"],
            ["Self-assessment deadline",  "Employee",               "48h / 24h / at deadline"],
            ["Self-assessment submitted", "Employee's manager",     "Immediately after submission"],
            ["Manager scoring deadline",  "Manager",                "48h / 24h / at deadline"],
            ["Assessment finalized",      "Employee",               "HR closes or bulk-finalizes"],
            ["Weekly digest",             "Opt-in, all roles",      "Monday 9am tenant-local"],
            ["System alerts",             "Admins only",            "Cycle failures, export failures"],
        ],
        col_widths_inches=[2.0, 2.0, 2.2],
    )
    add_heading(doc, "9.2 What you cannot turn off", level=2)
    add_bullet(doc, "Security-critical emails — password changes, MFA enrolment, new-device sign-in.")
    add_bullet(doc, "Account deactivation notices.")
    add_bullet(doc, "Break-glass access alerts (Super Admin after-hours use).")


def ch_10_journey(doc):
    add_heading(doc, "10. End-to-End Journey", level=1)
    add_para(
        doc,
        "How a single review cycle flows end-to-end across all roles. Use this as the canonical "
        "reference when onboarding a new admin or explaining the platform to a stakeholder.",
    )
    steps = [
        ("1",  "HR Admin creates cycle (draft) — picks framework, roster, deadlines"),
        ("2",  "HR Admin opens cycle — employees get email + see cycle on dashboard"),
        ("3",  "Employee fills self-assessment — auto-saves every 30s; uploads artifacts; submits → self_submitted"),
        ("4",  "Manager gets email — opens /team, filters Pending, reads self-score + artifacts, sets manager score + rationale, submits → manager_scored"),
        ("5",  "HR Admin monitors /hr dashboard — completion donuts, identifies stragglers"),
        ("6",  "HR Admin locks cycle after manager deadline — no new edits"),
        ("7",  "HR Admin clicks Bulk Finalize — composites computed; status → composite_computed"),
        ("8",  "HR Admin exports CSV for appraisal — RFC 4180, allowlist-filtered, audited"),
        ("9",  "HR Admin closes cycle — status → closed → read-only forever"),
        ("10", "Employee sees finalized scores on /scorecard — trend line updates"),
        ("11", "AI Champion reviews distributions via /hr/reports — feeds next-quarter training"),
    ]
    add_table(
        doc,
        ["Step", "What happens"],
        [list(s) for s in steps],
        col_widths_inches=[0.5, 5.8],
    )
    add_heading(doc, "Typical durations at Qualtech", level=2)
    add_table(
        doc,
        ["Phase", "Typical duration"],
        [
            ["Draft (HR setup)",              "2–3 business days"],
            ["Open for self-assessment",      "2 weeks"],
            ["Open for manager scoring",      "1 week"],
            ["Locked + finalize",             "1 day"],
            ["Appraisal integration + close", "1 day"],
            ["Total per cycle",               "~4 weeks"],
        ],
        col_widths_inches=[3.0, 3.3],
    )


def ch_11_security(doc):
    add_heading(doc, "11. Security, Privacy, and Your Data", level=1)

    add_heading(doc, "11.1 Authentication", level=2)
    add_bullet(doc, "Passwords stored as bcrypt with per-password salts. No recovery — only rotation.")
    add_bullet(doc, "Access tokens: 15-minute JWT expiry, silently refreshed by a rotating refresh token.")
    add_bullet(doc, "Tokens are HTTP-only cookies — JavaScript cannot read them (XSS defense).")
    add_bullet(doc, "SSO tokens verified against tenant-configured issuer AND audience.")

    add_heading(doc, "11.2 Multi-tenancy", level=2)
    add_bullet(doc, "Every DB row tagged with org_id. Every query filters by org_id.")
    add_bullet(doc, "Enforced at two layers: application guards + Postgres Row-Level Security (RLS).")
    add_bullet(doc, "Cross-tenant lookups return 404, not 403 — prevents existence leaks.")

    add_heading(doc, "11.3 Data encryption", level=2)
    add_bullet(doc, "In transit: TLS 1.3 on every public endpoint + HSTS.")
    add_bullet(doc, "At rest: RDS + S3 with AWS KMS-managed keys.")
    add_bullet(doc, "Field-level: email and phone use pgcrypto in addition to disk encryption.")
    add_bullet(doc, "Secrets in AWS Secrets Manager, never in .env files in prod images.")

    add_heading(doc, "11.4 Audit log", level=2)
    add_bullet(doc, "Every write on assessments, cycles, users, frameworks, artifacts → audit entry.")
    add_bullet(doc, "Actor, timestamp, old/new values, request ID.")
    add_bullet(doc, "Append-only — service role has no UPDATE/DELETE grants.")
    add_bullet(doc, "Retention: 7 years, per SOC 2 Type II.")

    add_heading(doc, "11.5 AI and PII", level=2)
    add_bullet(doc, "Name, email, and directly-identifying text stripped before every Claude API call.")
    add_bullet(doc, "AI sees pseudonymous employee IDs + scores + evidence — never identifiers.")
    add_bullet(doc, "AI outputs validated against schema before DB write.")
    add_bullet(doc, "Zero-data-retention Anthropic agreement — no training on your data.")

    add_heading(doc, "11.6 Rate limits + CSRF", level=2)
    add_bullet(doc, "Auth endpoints (login, refresh, accept-invite, SSO exchange) rate-limited at ~10/min/IP.")
    add_bullet(doc, "ThrottlerGuard runs before JWT verification — expensive crypto gated by cheap check.")
    add_bullet(doc, "All state-changing BFF POSTs validate Origin header against APP_ORIGIN_ALLOWLIST.")
    add_bullet(doc, "SameSite=Lax cookies handle most CSRF; Origin check handles the residual top-level form POST vector.")

    add_heading(doc, "11.7 Your rights under DPDP Act 2023", level=2)
    add_bullet(doc, "Right to access — request your data via HR.")
    add_bullet(doc, "Right to correction — name typos, role errors — HR can fix.")
    add_bullet(doc, "Right to erasure — cascading soft-delete + PII scrub within 30 days.")
    add_bullet(doc, "Consent records — invite-accept is logged with timestamp + privacy-notice version.")


def ch_12_phase_2(doc):
    add_heading(doc, "12. What's coming in Phase 2", level=1)
    add_para(doc, "Phase 2 — AI Intelligence — starts 2026-06-08 after the May-2026 appraisal cycle.")
    add_table(
        doc,
        ["Feature", "Who benefits", "Replaces / Adds"],
        [
            ["Claude-powered artifact analysis", "Managers", "'AI has read the 5 artifacts and highlighted 3' — cuts review time"],
            ["AI-suggested scores + confidence", "Managers", "Fills the Phase-1 AiSuggestionBadge placeholder"],
            ["Peer feedback",                    "Employees", "3–5 peers per cycle; anonymous aggregation"],
            ["HRMS integration",                 "HR Admin",  "Bi-directional sync with your HRIS"],
            ["Mobile app (React Native Expo)",   "All roles", "Scorecards + approvals on iOS/Android"],
            ["Bias detection",                   "AI Champion", "Manager-AI delta heat-map, role-group bias"],
            ["Leadership dashboard",             "Leadership", "Org-wide capability telescope"],
            ["In-app notifications",             "All",        "Complements email"],
        ],
        col_widths_inches=[2.2, 1.4, 2.7],
    )


def ch_13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting", level=1)
    items = [
        ("I can't sign in", [
            "Check spam for the invite — sender is no-reply@skillforge.qualtech.ai.",
            "Expired invite? Ask HR to re-invite (old links stay dead).",
            "SSO loop back to login? Usually an audience-mismatch on IdP side — screenshot URL (not tokens) and send to support.",
            "Password reset stuck? Links expire in 1 hour. Request again.",
            "MFA lost? Use a recovery code. If none, Super Admin can reset after identity check.",
        ]),
        ("My self-assessment won't submit", [
            "Every dimension filled? Submit is disabled until all are scored. Red dots in left nav show missing ones.",
            "Network offline? Auto-save shows 'Saving failed — will retry.' Don't close the tab.",
            "Session expired after 15 min? Re-login in another tab, come back, Save draft.",
        ]),
        ("I can't see someone I think I manage", [
            "Check /users shows them with you as their manager (HR Admin can check).",
            "Check the cycle has them on the roster.",
            "If both look right, file a ticket.",
        ]),
        ("The CSV doesn't import", [
            "Correct export template selected?",
            "UTF-8 with BOM — our default. If your tool chokes on BOM, config change needed.",
            "Case-sensitive column name match to your import spec.",
        ]),
        ("Artifact upload fails", [
            "File >25 MB? Hard limit.",
            "Wrong type? PDF/PNG/JPG/MP4/DOCX/XLSX/TXT only.",
            "'Upload URL expired' after 15 min? Re-drop the file — fresh URL issues.",
        ]),
    ]
    for title, bullets in items:
        add_heading(doc, title, level=2)
        for b in bullets:
            add_bullet(doc, b)
    add_callout(
        doc,
        "Always include the request-ID",
        "Every API response carries a request-ID header. When filing a ticket, paste it — fastest "
        "way for support to trace your exact call in the audit log.",
        color=BLUE,
        bg=LIGHT_BG,
    )


def ch_14_faq(doc):
    add_heading(doc, "14. FAQ", level=1)
    faq = [
        ("Can I see how my manager scored me before it's finalized?",
         "No. While status is manager_in_progress, scores are hidden. After composite_computed or finalized, you see them on /scorecard."),
        ("Can I re-submit my self-assessment if I change my mind?",
         "Not after submission. Ask your manager to capture the correction in their rationale, or ask HR to reopen if the cycle is still open."),
        ("Does SkillForge look at my Slack / email / calendar?",
         "No. SkillForge only sees what you type into it and artifacts you upload. Phase 2 may add opt-in integrations."),
        ("Who can see my free-text comments?",
         "Your manager, HR Admin, AI Champion, and Super Admin. Not Leadership. Not peers. Not other employees."),
        ("Can Claude see my name?",
         "No. Names and emails are stripped before any AI call. Claude sees pseudonymous IDs, scores, and evidence."),
        ("What happens if the platform goes down mid-assessment?",
         "Your last auto-save is safe. Refresh when the platform is back; Resume draft picks up where you left off."),
        ("Can I export my personal scorecard?",
         "Phase 2 adds a PDF download. Today, ask HR for an export."),
        ("Can a manager delete my submission?",
         "No. Managers can't delete. HR Admin or Super Admin can reopen for correction — reopens are audited."),
        ("I'm a manager and also an employee. How does that work?",
         "Cleanly. You self-assess in /assessments and score reports in /team. Your own manager scores you in their /team."),
        ("What browsers are supported?",
         "Latest two versions of Chrome, Edge, Firefox, Safari. Mobile browsers work; the polish ships with the Phase-2 mobile app."),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Q. ")
        set_run(r, size=10, color=NAVY, bold=True)
        r2 = p.add_run(q)
        set_run(r2, size=10, color=DARK, bold=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r3 = p2.add_run("A. ")
        set_run(r3, size=10, color=BLUE, bold=True)
        r4 = p2.add_run(a)
        set_run(r4, size=10, color=DARK)


def ch_15_glossary(doc):
    add_heading(doc, "15. Glossary", level=1)
    terms = [
        ("Assessment", "A single (cycle × employee) record with self-score, manager-score, AI suggestion, composite, artifacts."),
        ("Artifact", "A file attached to an assessment dimension as evidence."),
        ("Audit log", "Append-only record of every data-changing action in the tenant."),
        ("BFF", "Backend-For-Frontend. Next.js route handlers between browser and the API service."),
        ("Break-glass", "Elevated access used only in emergencies. Always audited."),
        ("Competency framework", "Definition of what dimensions to measure and how."),
        ("Composite score", "Single number per assessment — dimension-weight-weighted average of manager scores."),
        ("CSRF", "Cross-Site Request Forgery. SkillForge defends with Origin-header checks."),
        ("Cycle", "A named review round — e.g. 'H1 2026 AI Capability Review'."),
        ("Dimension", "A unit of measurement inside a framework — e.g. Prompting, Evaluation."),
        ("HRIS", "Human Resources Information System — your org's master HR database."),
        ("JWT", "JSON Web Token — the signed token that proves who you are on each request."),
        ("Manager override", "A manager choosing a different score than the AI suggestion. Rationale required."),
        ("Maturity level", "A rubric tier, 1–5, per dimension."),
        ("MFA", "Multi-Factor Authentication."),
        ("OIDC", "OpenID Connect. The SSO protocol SkillForge uses."),
        ("PII", "Personally Identifiable Information. Stripped before AI calls."),
        ("Refresh token", "Long-lived token that quietly replaces expired access tokens. Rotated on each use."),
        ("RLS", "Row-Level Security. Postgres feature enforcing tenant isolation at the engine level."),
        ("Scope claim", "A claim inside a signed short-lived URL that pins it to one action (upload/download/invite)."),
        ("SSO", "Single Sign-On."),
        ("Target maturity", "The level a role is expected to reach per dimension."),
        ("Tenant", "One customer organization's isolated slice of the platform."),
    ]
    add_table(
        doc,
        ["Term", "Meaning"],
        [list(t) for t in terms],
        col_widths_inches=[1.8, 4.5],
    )


def ch_16_support(doc):
    add_heading(doc, "16. Support and contact", level=1)
    add_table(
        doc,
        ["Channel", "Use for"],
        [
            ["skillforge-support@qualtech.ai",              "General questions, feature requests"],
            ["skillforge-security@qualtech.ai",             "Security incidents, lost MFA, break-glass"],
            ["#skillforge-support (Slack)",                 "Quick questions, peer help"],
            ["https://status.skillforge.qualtech.ai",       "Known incidents, planned maintenance"],
            ["https://skillforge.qualtech.ai/docs",         "This guide, admin runbooks, API reference"],
        ],
        col_widths_inches=[2.8, 3.5],
    )

    add_heading(doc, "Response targets", level=2)
    add_table(
        doc,
        ["Priority", "First response", "Resolution target"],
        [
            ["P0 — Can't sign in / data loss",  "1 business hour",  "4 business hours"],
            ["P1 — Cycle operations blocked",    "4 business hours", "1 business day"],
            ["P2 — UI bug, feature question",    "1 business day",   "5 business days"],
            ["P3 — Feature request",             "5 business days",  "Assessed next sprint"],
        ],
        col_widths_inches=[2.5, 1.8, 2.0],
    )

    add_para(doc, "", space_before=24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "End of User Guide v1.0 · SkillForge AI · Qualtech · "
        f"prepared by {DOC_AUTHOR} for {DOC_FOR}"
    )
    set_run(r, size=9, color=MEDIUM, italic=True)


# ═══════════════════════════════════════════════════════════════════════
#  Build
# ═══════════════════════════════════════════════════════════════════════

def build() -> Path:
    doc = Document()

    # Base style — Arial 10 dark body
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    # east-asian + complex-script fonts
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Arial")

    # One section — Letter, 1" margins, header + footer
    section = doc.sections[0]
    set_margins(section, 1.0, 1.0, 1.0, 1.0)
    set_header_footer(section)

    # Build
    write_cover(doc)
    write_version_history(doc)
    write_toc(doc)

    ch_1_welcome(doc)
    ch_2_getting_started(doc)
    page_break(doc)
    ch_3_employees(doc)
    page_break(doc)
    ch_4_managers(doc)
    page_break(doc)
    ch_5_hr_admins(doc)
    page_break(doc)
    ch_6_ai_champions(doc)
    ch_7_leadership(doc)
    page_break(doc)
    ch_8_super_admins(doc)
    page_break(doc)
    ch_9_notifications(doc)
    ch_10_journey(doc)
    page_break(doc)
    ch_11_security(doc)
    page_break(doc)
    ch_12_phase_2(doc)
    ch_13_troubleshooting(doc)
    page_break(doc)
    ch_14_faq(doc)
    ch_15_glossary(doc)
    page_break(doc)
    ch_16_support(doc)

    # Save
    doc.save(OUT_PATH)

    # Post-process: normalize rare glyphs that can render as tofu in older
    # Windows Arial builds. Arial ships strong coverage of ●, —, –, ·, …, ×, §
    # (we keep those), but can fall back or tofu on U+2192 →, U+2265 ≥, and
    # U+03C3 σ depending on installation. Substitute those three into safe
    # ASCII equivalents after the DOCX is built — applied to every .xml part
    # (document, headers, footers, core/app props) so TOC + metadata are covered.
    _sanitize_docx_for_word(OUT_PATH)

    return OUT_PATH


# Char → ASCII-safe replacement. Keep the visual meaning; don't over-reach.
WORD_SAFE_SUBS = {
    "\u2192": "->",     # RIGHTWARDS ARROW
    "\u2265": ">=",     # GREATER-THAN OR EQUAL TO
    "\u03c3": "sigma",  # GREEK SMALL LETTER SIGMA
}


def _sanitize_docx_for_word(path: Path) -> None:
    """Replace tofu-prone glyphs in every XML part of a built DOCX in-place."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(path, "r") as zin, \
                zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                    text = data.decode("utf-8")
                    for bad, good in WORD_SAFE_SUBS.items():
                        text = text.replace(bad, good)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


if __name__ == "__main__":
    out = build()
    size_kb = os.path.getsize(out) / 1024
    print(f"✓ Wrote {out}")
    print(f"  {size_kb:,.1f} KB")
