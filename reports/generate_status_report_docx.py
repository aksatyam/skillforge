#!/usr/bin/env python3
"""
SkillForge AI — v1.0 Status Report DOCX generator.

Produces a branded, enterprise-grade status report following the 18-section
master template from ~/.claude/CLAUDE.md. Output is aimed at Qualtech
stakeholders (sponsors, Tech Lead, Security Lead, HR Lead) in the final
stretch before the 2026-06-01 go-live.

Brand standards (per global CLAUDE.md):
  - Navy #1B3A5C (headings, header rows, cover)
  - Blue #2E75B6 (accent, links, commit hashes)
  - Green #27AE60 (completed / on-track)
  - Orange #E67E22 (at-risk / in-progress)
  - Red #E74C3C (delayed / classification banner)
  - Arial body, Consolas code
  - US Letter, 1-inch margins, alt-row #F8F9FA
  - Status badges: ● (U+25CF) with colored text, no emoji circles

Attribution: Qualtech only — no third-party prepared-by line.

Run:
  python3 reports/generate_status_report_docx.py

Output:
  reports/SkillForge-Status-Report-v1.0.docx
"""

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand palette ─────────────────────────────────────────────────────

NAVY     = RGBColor(0x1B, 0x3A, 0x5C)
BLUE     = RGBColor(0x2E, 0x75, 0xB6)
GREEN    = RGBColor(0x27, 0xAE, 0x60)
ORANGE   = RGBColor(0xE6, 0x7E, 0x22)
RED      = RGBColor(0xE7, 0x4C, 0x3C)
DARK     = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM   = RGBColor(0x7F, 0x8C, 0x8D)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

ALT_ROW      = "F8F9FA"
HEADER_BG    = "1B3A5C"
SUCCESS_BG   = "E8F5E9"
WARN_BG      = "FFF8E1"
DANGER_BG    = "FFEBEE"
LIGHT_BG     = "F0F4F8"

# ── Metadata ──────────────────────────────────────────────────────────

DOC_TITLE     = "SkillForge AI — Project Status Report"
DOC_SUBTITLE  = "Phase-1 Hyper-MVP · Pre-Cutover Readiness Review"
DOC_VERSION   = "1.0"
DOC_DATE      = date(2026, 4, 19)
DOC_OWNER     = "Qualtech"
DOC_AUTHOR    = "Ashish Kumar Satyam"
DOC_FOR       = "Qualtech Executive Sponsors"
DOC_CLASS     = "Internal / Confidential"
DOC_ID        = "QTECH-SKILLFORGE-STATUS-V1.0-2026-04-19"

OUT_PATH = Path(__file__).parent / "SkillForge-Status-Report-v1.0.docx"


# ═══════════════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════════════

def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run(run, *, font="Arial", size=10, color=DARK, bold=False, italic=False):
    run.font.name = font
    r = run._element
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def add_para(doc, text, *, size=10, color=DARK, bold=False, italic=False,
             align=None, space_after=None, space_before=None, font="Arial"):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run(run, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, *, level=1):
    """H1 = 18pt navy bold, H2 = 14pt navy bold, H3 = 12pt blue bold."""
    size = {1: 18, 2: 14, 3: 12}.get(level, 11)
    color = NAVY if level <= 2 else BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text, *, size=10, color=DARK):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r, size=size, color=color)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def add_table(doc, headers, rows, *, widths=None, header_bg=HEADER_BG,
              header_color=WHITE, alt=True, font_size=9, first_col_bold=False):
    """Branded table with navy header + alt-row shading."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False

    # Header
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        shade(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        set_run(r, size=font_size + 1, color=header_color, bold=True)
        if widths and j < len(widths):
            cell.width = Inches(widths[j])

    # Body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if alt and i % 2 == 1:
                shade(cell, ALT_ROW)
            if widths and j < len(widths):
                cell.width = Inches(widths[j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Status-badge convention: cells starting with "● " get green/orange/red coloring
            if isinstance(val, str) and val.startswith("● "):
                label = val[2:].lower()
                if any(k in label for k in ("complete", "done", "on track", "pass", "green", "ok", "delivered")):
                    dot_color = GREEN; bg = SUCCESS_BG
                elif any(k in label for k in ("progress", "risk", "pending", "yellow", "in flight", "prep", "scheduled")):
                    dot_color = ORANGE; bg = WARN_BG
                elif any(k in label for k in ("delay", "block", "fail", "red", "overdue")):
                    dot_color = RED; bg = DANGER_BG
                else:
                    dot_color = GREEN; bg = SUCCESS_BG
                shade(cell, bg)
                r = p.add_run("● ")
                set_run(r, size=font_size, color=dot_color, bold=True)
                r2 = p.add_run(val[2:])
                set_run(r2, size=font_size, color=DARK)
            else:
                r = p.add_run(str(val))
                bold = first_col_bold and j == 0
                # Commit hashes in Consolas blue
                if isinstance(val, str) and len(val) == 7 and all(c in "0123456789abcdef" for c in val):
                    set_run(r, font="Consolas", size=font_size, color=BLUE)
                else:
                    set_run(r, size=font_size, color=DARK, bold=bold)

    return tbl


def add_metric_card(doc, cards):
    """4-card metric strip in a 1×4 table with navy headers."""
    tbl = doc.add_table(rows=2, cols=len(cards))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for j, (label, value, color) in enumerate(cards):
        # Label row
        c1 = tbl.cell(0, j)
        shade(c1, HEADER_BG)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.width = Inches(1.6)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        set_run(r, size=9, color=WHITE, bold=True)
        # Value row
        c2 = tbl.cell(1, j)
        shade(c2, ALT_ROW)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.width = Inches(1.6)
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(value)
        set_run(r, size=18, color=color, bold=True)


def add_callout(doc, kind, text):
    """Colored callout block. kind ∈ {info, warn, danger, success}."""
    color_map = {
        "info":    (BLUE,   LIGHT_BG,   "NOTE"),
        "warn":    (ORANGE, WARN_BG,    "RISK"),
        "danger":  (RED,    DANGER_BG,  "BLOCKER"),
        "success": (GREEN,  SUCCESS_BG, "OK"),
    }
    c, bg, label = color_map[kind]
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label}  ")
    set_run(r, size=9, color=c, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=10, color=DARK)


# ═══════════════════════════════════════════════════════════════════════
#  Header + footer
# ═══════════════════════════════════════════════════════════════════════

def set_header_footer(section):
    # Header
    hdr = section.header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"{DOC_TITLE}  ·  v{DOC_VERSION}")
    set_run(hr, size=9, color=MEDIUM, italic=True)

    # Footer with "Page X of Y"
    ftr = section.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rpr.append(rfonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rpr.append(sz)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "7F8C8D"); rpr.append(col)
        r.append(rpr)
        t = OxmlElement("w:t"); t.text = "1"; r.append(t)
        fld.append(r)
        return fld

    # Prefix
    r = fp.add_run("Page ")
    set_run(r, size=9, color=MEDIUM)
    fp._p.append(_field("PAGE"))
    r = fp.add_run(" of ")
    set_run(r, size=9, color=MEDIUM)
    fp._p.append(_field("NUMPAGES"))
    r = fp.add_run(f"   ·   {DOC_ID}   ·   {DOC_CLASS}")
    set_run(r, size=9, color=MEDIUM, italic=True)


# ═══════════════════════════════════════════════════════════════════════
#  Cover page
# ═══════════════════════════════════════════════════════════════════════

def write_cover(doc):
    for _ in range(3):
        add_para(doc, "", size=10)

    # Company
    add_para(doc, DOC_OWNER, size=16, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "AI-Powered Employee Skill Assessment Platform",
             size=10, color=MEDIUM, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    # Title
    add_para(doc, "SkillForge AI", size=32, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Project Status Report", size=22, color=BLUE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, DOC_SUBTITLE, size=12, color=DARK, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Classification banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(3.5)
    shade(cell, "FFEBEE")
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f"  {DOC_CLASS}  ")
    set_run(r, size=11, color=RED, bold=True)

    add_para(doc, "", space_before=30)

    # Metadata table (Section 1 content on cover for quick scan)
    rows = [
        ["Document ID", DOC_ID],
        ["Version", DOC_VERSION],
        ["Report Date", DOC_DATE.isoformat()],
        ["Reporting Period", "2026-04-18 → 2026-04-19  (Sprints 0–6 + hardening)"],
        ["Owner", DOC_OWNER],
        ["Author", DOC_AUTHOR],
        ["Audience", DOC_FOR],
        ["Classification", DOC_CLASS],
        ["Phase", "Phase 1 (Hyper-MVP) — pre-cutover"],
        ["Next Milestone", "Production deploy  ·  2026-05-29"],
    ]
    meta = doc.add_table(rows=len(rows), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    for i, (k, v) in enumerate(rows):
        c1 = meta.cell(i, 0)
        c2 = meta.cell(i, 1)
        c1.width = Inches(1.8)
        c2.width = Inches(4.2)
        if i % 2 == 1:
            shade(c1, ALT_ROW); shade(c2, ALT_ROW)
        for cc in (c1, c2):
            cc.paragraphs[0].paragraph_format.space_before = Pt(3)
            cc.paragraphs[0].paragraph_format.space_after = Pt(3)
        r1 = c1.paragraphs[0].add_run(k)
        set_run(r1, size=10, color=NAVY, bold=True)
        r2 = c2.paragraphs[0].add_run(v)
        set_run(r2, size=10, color=DARK)

    add_para(doc, "", space_before=40)

    # "Prepared for" navy block (Qualtech-only)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, HEADER_BG)
    for text, bold in [
        ("Prepared for", False),
        (DOC_FOR, True),
        (f"by {DOC_AUTHOR}", False),
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        set_run(r, size=11, color=WHITE, bold=bold)
    first = cell.paragraphs[0]
    if first.text == "":
        first._element.getparent().remove(first._element)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════
#  Sections (18)
# ═══════════════════════════════════════════════════════════════════════

def s1_document_information(doc):
    add_heading(doc, "1. Document Information", level=1)
    add_para(doc,
        "This report summarizes the development, security, and readiness posture of SkillForge AI "
        "as of 2026-04-19, heading into the 2026-06-01 Qualtech appraisal-cycle go-live. It consolidates "
        "sprint delivery evidence, security audit outcomes, infrastructure status, outstanding risks, and "
        "the UAT/cutover plan for executive review."
    )
    add_table(doc,
        ["Field", "Value"],
        [
            ["Document Title", DOC_TITLE],
            ["Document ID", DOC_ID],
            ["Version", DOC_VERSION],
            ["Issue Date", DOC_DATE.isoformat()],
            ["Owner", DOC_OWNER],
            ["Author", DOC_AUTHOR],
            ["Audience", DOC_FOR],
            ["Classification", DOC_CLASS],
            ["Distribution", "Executive sponsors · Tech Lead · Security Lead · HR Lead · PMO"],
            ["Review Cadence", "Weekly until go-live; monthly thereafter"],
            ["Supersedes", "STATUS.md (2026-04-18)"],
        ],
        widths=[1.8, 4.7],
    )


def s2_version_history(doc):
    add_heading(doc, "2. Version History", level=1)
    add_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["0.1", "2026-04-18", DOC_AUTHOR,
             "Initial status snapshot (text-only STATUS.md) — Sprints 0–3 closed."],
            ["0.2", "2026-04-19 AM", DOC_AUTHOR,
             "Sprint 4 + 5 + 6 deltas integrated. Test count 127/127 green."],
            ["1.0", "2026-04-19 PM", DOC_AUTHOR,
             "First formal stakeholder report: 18 sections, brand-compliant DOCX, "
             "post-Sprint-6 security audit embedded, cutover plan finalized."],
        ],
        widths=[0.7, 1.1, 1.6, 3.1],
    )


def s3_executive_summary(doc):
    add_heading(doc, "3. Executive Summary", level=1)
    add_para(doc,
        "SkillForge AI has completed its full Phase-1 Hyper-MVP scope — seven sprints (Sprint 0 through "
        "Sprint 6) plus two hardening passes — in a two-day wall-clock build window ending 2026-04-19. "
        "All committed P0 features from BUILD_PLAN §5 are in main. The test suite stands at 127 of 127 "
        "assertions green across 13 spec files, spanning tenant isolation, assessment scoring, artifact "
        "storage, export templates, notifications, statistics, and authentication.",
        space_after=6,
    )
    add_para(doc,
        "The platform has zero open Critical or High security findings. A comprehensive post-Sprint-6 "
        "audit surfaced and closed three Critical (C1–C3), six High (H1, H2, H4, H5, H6), two Medium, "
        "and one Low defect — every one landed in commit 4b23edb before cutover preparation. "
        "Artifact tokens migrated from no-expiry HMAC to jose-signed JWTs; OIDC aud validation, "
        "Origin-allowlist CSRF defense, and named-bucket rate limiting are all in place. Two ADRs "
        "(ADR-011 S3 storage, ADR-012 export-template allowlist) formalize the new surface.",
        space_after=6,
    )
    add_para(doc,
        "Production cutover is scheduled for Friday 2026-05-29 with go-live at 09:00 IST on Monday "
        "2026-06-01. The 43-day freeze window between now and deploy will be consumed by (a) UAT with "
        "20 pilot users from 2026-05-18 to 05-27, (b) screenshot capture and documentation finalization, "
        "and (c) final DR/rollback rehearsal per the deployment runbook. The primary residual risk is "
        "UAT fragility — features built at this cadence have not yet been exercised by real users — and "
        "is mitigated by the 9-day UAT window and the 72-hour post-launch war-room.",
    )


def s4_dashboard_view(doc):
    add_heading(doc, "4. Dashboard View", level=1)

    # 4 metric cards
    add_metric_card(doc, [
        ("COMMITS",      "13",      NAVY),
        ("TESTS GREEN",  "127/127", GREEN),
        ("OPEN C/HIGH",  "0",       GREEN),
        ("DAYS TO GO-LIVE", "43",   ORANGE),
    ])
    add_para(doc, "", space_before=8)

    add_heading(doc, "Workstream status", level=2)
    add_table(doc,
        ["Workstream", "Owner", "Status", "Notes"],
        [
            ["Core platform (Sprints 0–6)", "Tech Lead",  "● Delivered",
             "All 7 sprints committed; 127/127 tests green"],
            ["Security hardening (post-S6)", "Security Lead", "● Delivered",
             "C1/C2/C3 + H1–H6 + L2 + M1/M6 closed in 4b23edb"],
            ["User documentation",          "Docs",       "● Delivered",
             "Master guide + 6 one-pagers + branded DOCX (1f07c01)"],
            ["Screenshot capture",          "UX",         "● In Progress",
             "14 shots scoped; targeting 2026-05-10 → 05-17"],
            ["Staging environment",         "DevOps",     "● In Progress",
             "SSO config + seed data in prep; smoke pass 2026-05-15"],
            ["UAT (20 pilot users)",        "HR Lead",    "● Scheduled",
             "Kickoff 2026-05-18; sign-off deadline 2026-05-28 EOD"],
            ["Production deploy",           "DevOps",     "● Scheduled",
             "2026-05-29 EOD; G1–G7 pre-flight gates per runbook"],
            ["Phase-2 design",              "Tech Lead",  "● Scheduled",
             "Kickoff 2026-06-08; AI intelligence sprint series"],
        ],
        widths=[2.2, 1.2, 1.3, 2.0],
    )


def s5_completed_work(doc):
    add_heading(doc, "5. Completed Work", level=1)
    add_para(doc,
        "Every entry below is in origin/main and ships as part of v1.0.0-hyper-mvp on 2026-05-29.",
        italic=True, color=MEDIUM, space_after=8,
    )
    add_table(doc,
        ["Feature", "Description", "Sprint", "Commit"],
        [
            ["Monorepo scaffold", "pnpm + Turborepo, Next.js + NestJS workspaces, 10 ADRs, CI pipelines", "S0", "5397154"],
            ["Auth + RBAC",       "bcrypt login, JWT + refresh rotation, invite-accept flow, /auth/me",   "S1", "5397154"],
            ["Competency Framework engine", "Maturity levels, role mappings, publish/archive",            "S1", "5397154"],
            ["User invite UI",    "HR admin invites + one-time token acceptance",                         "S1", "5397154"],
            ["RLS + prismaAdmin", "BYPASSRLS admin client, transactional auth, frontend route guard",     "S1+", "99bc173"],
            ["Self-assessment backend", "Draft save + submit; 0–5 scale × 0.01; 2000-char comments",      "S2", "2d98fff"],
            ["Artifact upload",   "jose-signed storage tokens (15m upload / 5m download); tenant-scoped", "S2", "2d98fff"],
            ["Manager roster",    "/team with filter pills + status-aware CTAs",                          "S2", "2d98fff"],
            ["BullMQ reminders",  "Email reminder subsystem for unsubmitted assessments",                 "S2", "2d98fff"],
            ["Manager scoring UI", "Rubric-weighted average + composite preview + rationale gate",        "S3", "35d9710"],
            ["CSV export",        "RFC 4180 + UTF-8 BOM; deterministic column order for appraisal sync",  "S3", "35d9710"],
            ["Cycle state machine", "Open → lock → bulk-finalize → close with audit trail",               "S3", "35d9710"],
            ["HR dashboard",      "KPI strip, cycle cards with donut progress, roster table",             "S3", "35d9710"],
            ["Employee scorecard", "Radar chart + histogram + cycle table at /scorecard",                 "S4", "d92637a"],
            ["Team overview",     "Dual-layer radar + per-person mini-cards + completion donut",          "S4", "d92637a"],
            ["HR reports",        "Capability heat-map + outlier list + training-need feed",              "S4", "d92637a"],
            ["HTML emails",       "Role-appropriate templates for invites + reminders + notifications",   "S5", "ad9f174"],
            ["Notification prefs", "Per-user opt-out for non-mandatory emails at /settings/notifications", "S5", "ad9f174"],
            ["httpOnly cookies + Swagger", "Session cookies moved off JS; Swagger restored on dev",       "S5", "ad9f174"],
            ["SSO (Keycloak OIDC)", "aud validation, state + nonce, bridge-secret enforcement",           "S6", "d1db56f"],
            ["S3 artifact storage", "Provider-swappable; jose-signed tokens carry orgId + scope + exp",   "S6", "d1db56f"],
            ["HR-editable CSV templates", "Default-deny column allowlist at /hr/templates",               "S6", "d1db56f"],
            ["Post-S6 security audit", "3 Critical + 6 High + 2 Medium + 1 Low closed",                   "S6+", "4b23edb"],
            ["Application user guide", "Master + 6 role one-pagers + branded 56 KB DOCX",                 "S6+", "1f07c01"],
        ],
        widths=[1.6, 3.0, 0.5, 0.9],
    )


def s6_work_in_progress(doc):
    add_heading(doc, "6. Work in Progress", level=1)
    add_table(doc,
        ["Item", "Description", "Owner", "Target"],
        [
            ["Screenshot capture",
             "14 PNGs per docs/user-guide/screenshots/README.md capture brief (Chrome 1440×900, DPR 2×, light mode, fake data)",
             "UX", "2026-05-17"],
            ["DOCX screenshot embed",
             "Extend generate_user_guide_docx.py with insert_image() pass once PNGs land",
             "Docs", "2026-05-19"],
            ["Staging environment",
             "Keycloak realm, seed data (2 orgs, 20 users, 1 active cycle), SSO end-to-end validation",
             "DevOps", "2026-05-15"],
            ["UAT intake forms",
             "Per-persona checklist + feedback form + sign-off template (building on docs/ops/UAT_CHECKLIST.md)",
             "HR Lead", "2026-05-16"],
            ["Runbook rehearsal",
             "Dry-run DEPLOYMENT_RUNBOOK.md G1–G7 gates + rollback in staging",
             "DevOps", "2026-05-26"],
            ["Phase-2 design spike",
             "AI artifact analyzer prompt chain + PII anonymizer + confidence-band rubric",
             "Tech Lead", "2026-06-01 (kickoff prep)"],
        ],
        widths=[1.7, 3.2, 1.0, 1.1],
    )


def s7_pending_items(doc):
    add_heading(doc, "7. Pending Items", level=1)
    add_table(doc,
        ["Item", "Description", "Priority", "Remarks"],
        [
            ["Per-role dashboard routing",
             "/leadership currently falls back to /dashboard with read-only KPI strip — full split deferred to Phase 2",
             "P2", "Documented in user guide Chapter 13"],
            ["Break-glass audit-log viewer",
             "Super Admin reads audit_log via DB direct; in-app viewer deferred to Phase 2",
             "P2", "Flow documented in SUPER_ADMIN guide"],
            ["Mobile app (React Native)",
             "Native apps deferred to Phase 2 post-mid-June 2026",
             "P2", "Web is fully responsive on handset"],
            ["AI suggestion pipeline",
             "AiSuggestionBadge is a placeholder until Claude chain lands",
             "P2", "Manager scoring is 100% human-driven in Phase 1"],
            ["Peer feedback module",
             "360° input + visibility rules",
             "P2", "Phase 2 Sprint 3"],
            ["HRMS bidirectional integration",
             "Two-way sync with target HRIS (TBD based on Qualtech choice)",
             "P2", "Phase 2 Sprint 4+"],
            ["Bias detection",
             "Statistical outlier analysis across managers with calibration dashboard",
             "P2", "Phase 2 Sprint 5"],
            ["Backfill of historical cycles",
             "Pre-2026 assessments import (optional)",
             "P3", "Only if HR requests"],
        ],
        widths=[1.8, 3.1, 0.7, 1.4],
    )


def s8_future_roadmap(doc):
    add_heading(doc, "8. Future Roadmap", level=1)
    add_table(doc,
        ["Phase", "Feature Cluster", "Timeline", "Effort"],
        [
            ["Phase 1", "Hyper-MVP (self + manager + HR flows, CSV export)", "Apr – Jun 2026", "7 sprints ✓"],
            ["Phase 1+", "Pre-cutover hardening + UAT + go-live", "Apr 19 – Jun 1 2026", "43 days"],
            ["Phase 2.1", "AI artifact analysis (Claude) + AI-suggested scores", "Jun – Jul 2026", "~6 weeks"],
            ["Phase 2.2", "Peer feedback (360°) + prompt library", "Jul – Aug 2026", "~4 weeks"],
            ["Phase 2.3", "Advanced analytics: heat-maps over time, training-ROI tracker", "Aug – Sep 2026", "~6 weeks"],
            ["Phase 2.4", "HRMS bidirectional integration + SSO extension", "Sep – Oct 2026", "~6 weeks"],
            ["Phase 2.5", "Mobile React Native app (Expo)", "Oct – Nov 2026", "~6 weeks"],
            ["Phase 2.6", "Bias detection + calibration dashboard (AI Champion)", "Nov – Dec 2026", "~4 weeks"],
            ["Phase 3",   "External SaaS resale: multi-tenant billing, provisioning, onboarding", "2027 H1", "TBD per sales demand"],
        ],
        widths=[0.9, 3.3, 1.6, 1.2],
    )


def s9_upcoming_deliveries(doc):
    add_heading(doc, "9. Upcoming Deliveries", level=1)
    add_table(doc,
        ["Version", "Scope", "Planned Release"],
        [
            ["v1.0.0-rc1",     "Staging deploy for UAT kickoff", "2026-05-17"],
            ["v1.0.0-rc2",     "UAT feedback fixes (if any)",    "2026-05-26"],
            ["v1.0.0-hyper-mvp", "Production deploy — tag + release notes", "2026-05-29 EOD"],
            ["v1.0.1",         "Post-launch patches (expected: minor copy + CSV format tweaks)", "2026-06-05 to 06-12"],
            ["v1.1.0",         "Phase 2.1 landing: AI artifact analysis + suggested scores",    "2026-07-20 (target)"],
            ["v1.2.0",         "Peer feedback + prompt library",                                 "2026-08-31 (target)"],
        ],
        widths=[1.3, 3.8, 1.7],
    )


def s10_risks_dependencies(doc):
    add_heading(doc, "10. Risks & Dependencies", level=1)
    add_table(doc,
        ["Risk", "Impact", "Likelihood", "Mitigation"],
        [
            ["UAT surfaces blocking functional bug",
             "Go-live slip",
             "Medium",
             "9-day UAT window + daily triage; rollback plan in DEPLOYMENT_RUNBOOK.md G7"],
            ["SSO/Keycloak integration hiccup with Qualtech IdP",
             "Sign-in outage",
             "Medium",
             "Dev-mode fallback to email+password; bridge-secret pre-validated at boot (H5)"],
            ["Data-model migration drift between dev/staging/prod",
             "Data loss / reject",
             "Low",
             "Prisma migrations are baseline-tracked; G2 pre-flight runs migration dry-run"],
            ["Operator error during cutover",
             "Corrupt data or downtime",
             "Low",
             "Runbook has G1–G7 checklist; two-person confirmation required at G4 (schema apply)"],
            ["AI-vendor lock-in (Claude)",
             "Future cost/flex",
             "Low",
             "AI layer isolated behind LangChain + prompts directory; swap-friendly"],
            ["HR Admin unfamiliarity with cycle lifecycle",
             "Support tickets on day 1",
             "Medium",
             "USER_GUIDE_HR_ADMIN + 90-min training session scheduled 2026-05-27"],
            ["Post-launch monitoring gap",
             "Slow incident detection",
             "Low",
             "72h war-room on-call roster + Sentry + Grafana alerts pre-configured"],
            ["Calendar: May 29 is a Friday deploy",
             "Weekend on-call",
             "Medium",
             "War-room extended to Monday midday; PagerDuty rotations confirmed"],
        ],
        widths=[1.9, 1.1, 0.9, 2.8],
    )


def s11_architecture_overview(doc):
    add_heading(doc, "11. Architecture Overview", level=1)
    add_para(doc,
        "Phase-1 consolidates the plan's 6-service target into a single NestJS service "
        "(assessment-service) for delivery speed. The service will be broken into its full "
        "microservice topology in Phase 2, when the AI + integration surface forces the split.",
        italic=True, color=MEDIUM, space_after=8,
    )

    add_heading(doc, "Tech stack", level=2)
    add_table(doc,
        ["Layer", "Choice", "Notes"],
        [
            ["Backend",    "NestJS (TypeScript)",         "Single service: assessment-service (13 modules)"],
            ["Frontend",   "Next.js 14+ App Router + Tailwind + shadcn/ui", "25+ route segments, role-gated AppShell"],
            ["Mobile",     "React Native (Expo)",         "Deferred — Phase 2"],
            ["Database",   "PostgreSQL 15+ (Prisma)",     "12 Prisma models; Row-Level Security + BYPASSRLS admin role"],
            ["Cache / Queue", "Redis + BullMQ",           "Session fallback + reminder worker"],
            ["Object storage", "S3-compatible (pluggable provider)", "ADR-011; jose-signed 15m upload / 5m download tokens"],
            ["AI", "Claude API (Phase 2)",                "LangChain; PII anonymizer on all calls; deferred to Phase 2"],
            ["Auth",       "Keycloak (OIDC) + bcrypt + JWT", "aud + iss validated; MFA mandatory for HR/AI/Super roles"],
            ["Infra",      "AWS ECS/EKS; Terraform",      "Single-region; multi-region in Phase 3"],
            ["Monitoring", "Grafana + Prometheus + Sentry", "Dashboards landing 2026-05-20"],
            ["CI/CD",      "GitHub Actions",              "typecheck / lint / test / build / security-scan on every PR"],
        ],
        widths=[1.2, 1.9, 3.5],
    )

    add_heading(doc, "Services and ports", level=2)
    add_table(doc,
        ["Component", "Port", "Role"],
        [
            ["assessment-service (NestJS)", "4000", "Core API: auth, users, cycles, assessments, artifacts, exports, reports"],
            ["apps/web (Next.js)",          "3000", "Unified UI for all 6 roles; BFF proxy routes in /api/session/**"],
            ["PostgreSQL",                  "5432", "Primary DB (skillforge); skillforge_admin role has BYPASSRLS"],
            ["Redis",                       "6379", "Sessions + BullMQ queue"],
            ["Keycloak (staging+)",         "8080", "OIDC provider; Qualtech realm"],
            ["S3 / MinIO (dev)",            "9000", "Artifact storage"],
        ],
        widths=[2.2, 0.8, 3.6],
    )

    add_heading(doc, "Data model (12 Prisma models)", level=2)
    add_para(doc,
        "Organization · User · CompetencyFramework · RoleMapping · AssessmentCycle · Assessment · "
        "Artifact · PeerReview · LearningRecommendation · PromptLibraryEntry · AuditLog · RefreshToken",
        size=10, color=DARK, italic=True, space_after=4,
    )
    add_para(doc,
        "Every non-catalog table carries org_id and is filtered by Postgres RLS. Write-path audit "
        "coverage is enforced via a NestJS AuditLogInterceptor using the prismaAdmin client.",
        space_after=4,
    )


def s12_security_overview(doc):
    add_heading(doc, "12. Security Overview", level=1)
    add_para(doc,
        "Five non-negotiable security postures are locked in as of commit 4b23edb:",
        space_after=6,
    )
    add_table(doc,
        ["#", "Control", "Status", "Evidence"],
        [
            ["1", "Tenant isolation (RLS + withTenant + token-carries-tenant)",
             "● On Track", "Every artifact read/write scoped by {id, orgId}; C1 closed"],
            ["2", "Artifact tokens — jose-signed JWT with exp",
             "● On Track", "Upload 15m / download 5m; rotation invalidates all tokens; C2 closed"],
            ["3", "Secret-length gates (JWT_SECRET ≥32, SSO_BRIDGE_SECRET ≥32)",
             "● On Track", "Boot-time assertion fails loud; C3 + H5 closed"],
            ["4", "Rate limiting (ThrottlerGuard first in APP_GUARD chain)",
             "● On Track", "Named buckets: default 120/min, short 10/min on credential paths; H1 closed"],
            ["5", "Origin-header CSRF defense on every BFF POST",
             "● On Track", "APP_ORIGIN_ALLOWLIST checked via checkSameOrigin(); H2 closed"],
            ["6", "OIDC aud + iss validation",
             "● On Track", "Client-aud replay vector closed; H4 closed"],
            ["7", "Upstream error sanitization on SSO callback",
             "● On Track", "Logs where/status/error/error_description only; H6 closed"],
            ["8", "Append-only audit log (no UPDATE/DELETE grants)",
             "● On Track", "Enforced at Postgres role level; interceptor uses prismaAdmin"],
            ["9", "PII stripped before AI calls (when Phase 2 lands)",
             "● Scheduled", "Anonymizer pattern locked in; hook blocks raw-PII prompt edits"],
        ],
        widths=[0.4, 2.7, 1.3, 2.2],
    )
    add_para(doc, "", space_before=6)

    add_heading(doc, "Compliance posture", level=2)
    add_table(doc,
        ["Framework", "Status", "Notes"],
        [
            ["OWASP ASVS Level 2", "● On Track", "All L2 requirements mapped; test coverage for guards + tenant-guard"],
            ["SOC 2 Type II",       "● In Progress", "Control evidence collection starts 2026-06-01; 6-month window"],
            ["DPDP Act 2023 (India)", "● On Track", "Data-residency AWS ap-south-1; consent records in Assessment model"],
            ["GDPR (future)",       "● Scheduled", "Required if external SaaS resale; Phase 3"],
        ],
        widths=[1.8, 1.3, 3.5],
    )


def s13_vapt_findings(doc):
    add_heading(doc, "13. Security Findings — Post-Sprint-6 Audit", level=1)
    add_para(doc,
        "A through-Sprint-6 audit on 2026-04-19 surfaced 3 Critical + 6 High + 2 Medium + 1 Low. "
        "All landed in commit 4b23edb before cutover prep. The project enters UAT with zero open "
        "Critical or High findings.",
        space_after=8,
    )
    add_table(doc,
        ["ID", "Severity", "Title", "Status"],
        [
            ["C1", "Critical", "Artifact paths scoped only by id, not by (id, orgId)",            "● Fixed"],
            ["C2", "Critical", "Storage tokens had no expiry (HMAC only)",                       "● Fixed"],
            ["C3", "Critical", "JWT_SECRET had a dev fallback; could reach prod",                "● Fixed"],
            ["H1", "High",     "Rate limiting not running before JWT verification on auth paths","● Fixed"],
            ["H2", "High",     "No Origin-allowlist CSRF defense on BFF POST routes",           "● Fixed"],
            ["H4", "High",     "OIDC aud claim accepted any value within realm",                "● Fixed"],
            ["H5", "High",     "SSO_BRIDGE_SECRET length not enforced in production",           "● Fixed"],
            ["H6", "High",     "SSO callback logged raw upstream errors (token-leak vector)",   "● Fixed"],
            ["M1", "Medium",   "Verbose error messages in dev leaking stack traces to client",  "● Fixed"],
            ["M6", "Medium",   "Missing rate-limit on /users PATCH role-change endpoint",       "● Fixed"],
            ["L2", "Low",      "Swagger UI accessible in non-dev (now dev-only)",               "● Fixed"],
        ],
        widths=[0.5, 0.9, 4.1, 1.0],
    )


def s14_functional_gaps(doc):
    add_heading(doc, "14. Functional Gaps Summary", level=1)
    add_para(doc,
        "Phase-1 intentionally defers the following feature clusters to Phase 2+ to protect the "
        "2026-06-01 go-live. None are blocking for Qualtech's internal May–June appraisal cycle.",
        space_after=8,
    )
    add_table(doc,
        ["Gap", "Phase", "Effort", "Focus"],
        [
            ["AI-suggested scores (Claude)",      "Phase 2.1", "~6 weeks", "Manager acceleration + calibration signal"],
            ["Peer feedback (360°)",              "Phase 2.2", "~4 weeks", "Richer scoring triangulation"],
            ["Prompt library",                    "Phase 2.2", "~2 weeks", "Reusable AI chain building blocks"],
            ["Advanced analytics",                "Phase 2.3", "~6 weeks", "Heat-map over time + training-ROI"],
            ["Mobile app (RN + Expo)",            "Phase 2.5", "~6 weeks", "Field access for managers on the go"],
            ["Bias detection + calibration",      "Phase 2.6", "~4 weeks", "AI Champion toolkit"],
            ["HRMS bidirectional integration",    "Phase 2.4", "~6 weeks", "Cut CSV export handoff"],
            ["Multi-tenant billing + provisioning", "Phase 3", "~8 weeks", "External SaaS resale readiness"],
            ["Leadership dashboard (dedicated)",  "Phase 2.3", "~2 weeks", "Org-wide capability telescope"],
            ["Break-glass audit-log viewer",      "Phase 2.4", "~1 week",  "In-app Super Admin access to audit trail"],
        ],
        widths=[2.7, 0.9, 0.9, 2.0],
    )


def s15_environment_access(doc):
    add_heading(doc, "15. Environment Access", level=1)
    add_table(doc,
        ["Environment", "URL", "Auth", "Purpose"],
        [
            ["Local dev",    "http://localhost:3000",                     "email+password seed users",   "Developer smoke + feature work"],
            ["Staging",      "https://staging.skillforge.qualtech.ai (target)", "Keycloak SSO + seed",         "UAT pilot 2026-05-18 → 05-27"],
            ["Production",   "https://skillforge.qualtech.ai",            "Keycloak SSO + MFA",          "Go-live 2026-06-01"],
            ["Keycloak",     "https://sso.qualtech.ai (target)",          "Admin console",               "Realm: qualtech; Clients: skillforge-web"],
            ["Postgres (prod)", "Managed RDS (ap-south-1)",               "IAM + passwordless",          "Daily PITR backups + 7-day retention"],
            ["S3 bucket (prod)", "s3://skillforge-artifacts-prod",        "IAM role",                    "Artifact storage; versioned + encrypted"],
            ["CI/CD",        "GitHub Actions",                            "GitHub OIDC → AWS",           "Typecheck, lint, test, build, security-scan"],
            ["Monitoring",   "Grafana + Sentry",                          "SSO",                         "Landing 2026-05-20"],
            ["Status page",  "https://status.skillforge.qualtech.ai",     "Public",                      "Incident communication"],
            ["Repo",         "github.com/aksatyam/skillforge (private)",  "SSH key required",            "Source of truth"],
        ],
        widths=[1.3, 2.4, 1.4, 1.8],
    )
    add_callout(doc, "info",
        "Credentials and secret rotation are managed via AWS Secrets Manager. "
        "No shared accounts. Super Admin MFA is mandatory and enforced at first privileged action.")


def s16_effort_summary(doc):
    add_heading(doc, "16. Effort Summary", level=1)
    add_metric_card(doc, [
        ("COMMITS",       "13",       NAVY),
        ("LINES ADDED",   "29,441",   GREEN),
        ("FILES TOUCHED", "231",      BLUE),
        ("TESTS",         "127",      GREEN),
    ])
    add_para(doc, "", space_before=8)
    add_table(doc,
        ["Metric", "Value", "Notes"],
        [
            ["Total commits",        "13",          "S0–S6 + 2 hardening + 2 docs passes"],
            ["Insertions",           "29,441",      "Includes generated Prisma client + lockfiles"],
            ["Deletions",            "839",         "Mostly refactors in hardening commits"],
            ["Unique files touched", "231",         "Counted across all 13 commits"],
            ["Source lines (TS)",    "11,003",      "105 .ts files in services/ + packages/"],
            ["Source lines (TSX)",   "6,262",       "29 .tsx files in apps/web"],
            ["Source lines (Python)", "1,822",      "Docs generators only"],
            ["Source lines (SQL)",   "414",         "Migrations"],
            ["Source lines (MD)",    "5,619",       "48 markdown files including docs + ADRs"],
            ["Tests (total)",        "127",         "13 spec files; 100% passing"],
            ["ADRs",                 "13",          "000 template + 012 most recent (export allowlist)"],
            ["Prisma models",        "12",          "Organization through RefreshToken"],
            ["Dependencies (top-level)", "27",      "Runtime + dev; audited weekly"],
            ["Development window",   "2 days (wall-clock)", "2026-04-18 → 2026-04-19; heavy skill + agent assistance"],
        ],
        widths=[2.0, 1.4, 3.5],
    )


def s17_commit_breakdown(doc):
    add_heading(doc, "17. Commit Type Breakdown", level=1)
    add_table(doc,
        ["Type", "Count", "Share", "Purpose"],
        [
            ["feat",  "5",  "38.5%", "Sprint feature deliveries (S2, S3, S4, S5, S6)"],
            ["fix",   "4",  "30.8%", "Post-S1 hardening, local-dev fixes, CSV bearer fix, post-S6 audit"],
            ["docs",  "3",  "23.1%", "E2E smoke evidence, skill refresh, user guide + status"],
            ["chore", "1",  "7.7%",  "Initial Sprint-0 scaffold commit"],
        ],
        widths=[0.9, 0.8, 0.9, 4.3],
    )
    add_para(doc, "", space_before=6)
    add_heading(doc, "Full commit log", level=2)
    add_table(doc,
        ["Hash", "Date", "Subject"],
        [
            ["1f07c01", "2026-04-19", "docs: add application user guide v1.0 + sync security-audit checklist"],
            ["4b23edb", "2026-04-19", "fix(post-sprint-6): security + test hardening audit (C1/C2/C3, H1-H6, L2, M1, M6)"],
            ["d1db56f", "2026-04-19", "feat(sprint-6): SSO (Keycloak OIDC) + S3 artifact storage + HR-editable CSV templates"],
            ["ad9f174", "2026-04-19", "feat(sprint-5): HTML emails + notification prefs + httpOnly cookies + Swagger restored"],
            ["d92637a", "2026-04-18", "feat(sprint-4): dashboards + reporting — employee scorecard, team overview, HR reports"],
            ["859195e", "2026-04-18", "docs: E2E smoke test evidence — composite=3.91 verified, append-only audit log confirmed"],
            ["848eb2a", "2026-04-18", "fix(local-dev): bring-up works end-to-end (8 issues resolved)"],
            ["3f30c14", "2026-04-18", "fix(hr): CSV download uses fetch+blob so bearer stays in Authorization header"],
            ["35d9710", "2026-04-18", "feat(sprint-3): Hyper-MVP close — manager scoring + CSV export + HR dashboard"],
            ["2d98fff", "2026-04-18", "feat(sprint-2): self-assessment + artifact + manager roster + email reminders"],
            ["2c2b2aa", "2026-04-18", "docs(claude): update skills with Sprint 1 actual patterns"],
            ["99bc173", "2026-04-18", "fix: resolve Critical/High issues from Sprint 1 validation"],
            ["5397154", "2026-04-18", "chore: initial commit — Sprint 0 scaffold + Sprint 1 features"],
        ],
        widths=[0.9, 1.1, 4.9],
    )


def s18_references(doc):
    add_heading(doc, "18. References & Links", level=1)

    add_heading(doc, "Repositories", level=2)
    add_table(doc,
        ["Name", "URL", "Role"],
        [
            ["skillforge (main)", "github.com/aksatyam/skillforge", "Source of truth — main branch"],
        ],
        widths=[1.6, 3.0, 2.1],
    )

    add_heading(doc, "In-tree documentation", level=2)
    add_table(doc,
        ["Document", "Path", "Purpose"],
        [
            ["BUILD_PLAN.md",        "/BUILD_PLAN.md",                       "Master plan (phases, sprints, P0/P1/P2)"],
            ["STATUS.md",            "/STATUS.md",                           "Living text-only status (pre-this report)"],
            ["OPEN_DECISIONS.md",    "/OPEN_DECISIONS.md",                   "Open architectural calls → ADRs"],
            ["User Guide (master)",  "/docs/USER_GUIDE.md",                  "16-chapter user guide"],
            ["User Guide (roles)",   "/docs/user-guide/",                    "6 persona one-pagers"],
            ["User Guide (DOCX)",    "/docs/SkillForge-User-Guide-v1.0.docx","Branded enterprise DOCX"],
            ["Deployment runbook",   "/docs/ops/DEPLOYMENT_RUNBOOK.md",      "G1–G7 gates + rollback"],
            ["UAT checklist",        "/docs/ops/UAT_CHECKLIST.md",           "7 scenarios + sign-off"],
            ["ADRs",                 "/docs/adr/",                           "13 decisions from ORM to export templates"],
            ["Sprint demos",         "/docs/SPRINT_*_DEMO.md",               "6 demo evidence docs (S1–S6)"],
            ["E2E smoke evidence",   "/docs/E2E_SMOKE_RESULT.md",            "Composite=3.91 + audit verification"],
        ],
        widths=[1.8, 2.5, 2.4],
    )

    add_heading(doc, "Operations + monitoring", level=2)
    add_table(doc,
        ["Service", "URL", "Owner"],
        [
            ["Production app",     "https://skillforge.qualtech.ai",              "DevOps"],
            ["Staging app",        "https://staging.skillforge.qualtech.ai",      "DevOps"],
            ["Status page",        "https://status.skillforge.qualtech.ai",       "DevOps on-call"],
            ["Sentry",             "sentry.qualtech.ai/projects/skillforge",      "Tech Lead"],
            ["Grafana",            "grafana.qualtech.ai/d/skillforge",            "DevOps"],
            ["GitHub Actions",     "github.com/aksatyam/skillforge/actions",      "Tech Lead"],
        ],
        widths=[1.8, 3.3, 1.6],
    )

    add_heading(doc, "Contacts", level=2)
    add_table(doc,
        ["Role", "Contact"],
        [
            ["Project Sponsor / CTO",   "cto@qualtech.ai"],
            ["Tech Lead",               "tech-lead@qualtech.ai"],
            ["Security Lead",           "skillforge-security@qualtech.ai"],
            ["HR Lead",                 "hr-lead@qualtech.ai"],
            ["Platform Support",        "skillforge-support@qualtech.ai"],
            ["On-call / Incidents",     "PagerDuty rotation: skillforge-prod"],
        ],
        widths=[2.2, 4.3],
    )

    add_para(doc, "", space_before=20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"— End of Status Report v{DOC_VERSION} · SkillForge AI · Qualtech · "
        f"issued {DOC_DATE.isoformat()} by {DOC_AUTHOR} —"
    )
    set_run(r, size=9, color=MEDIUM, italic=True)


# ═══════════════════════════════════════════════════════════════════════
#  Build + sanitize
# ═══════════════════════════════════════════════════════════════════════

WORD_SAFE_SUBS = {
    "\u2192": "->",
    "\u2265": ">=",
    "\u03c3": "sigma",
}


def _sanitize_docx(path: Path) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(path, "r") as zin, \
                zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                    s = data.decode("utf-8")
                    for bad, good in WORD_SAFE_SUBS.items():
                        s = s.replace(bad, good)
                    data = s.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def build() -> Path:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Arial")

    section = doc.sections[0]
    set_margins(section, 1.0, 1.0, 1.0, 1.0)
    set_header_footer(section)

    write_cover(doc)

    s1_document_information(doc); page_break(doc)
    s2_version_history(doc)
    s3_executive_summary(doc); page_break(doc)
    s4_dashboard_view(doc); page_break(doc)
    s5_completed_work(doc); page_break(doc)
    s6_work_in_progress(doc)
    s7_pending_items(doc); page_break(doc)
    s8_future_roadmap(doc)
    s9_upcoming_deliveries(doc); page_break(doc)
    s10_risks_dependencies(doc); page_break(doc)
    s11_architecture_overview(doc); page_break(doc)
    s12_security_overview(doc); page_break(doc)
    s13_vapt_findings(doc)
    s14_functional_gaps(doc); page_break(doc)
    s15_environment_access(doc); page_break(doc)
    s16_effort_summary(doc)
    s17_commit_breakdown(doc); page_break(doc)
    s18_references(doc)

    doc.save(OUT_PATH)
    _sanitize_docx(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    kb = os.path.getsize(out) / 1024
    print(f"[OK] Wrote {out}")
    print(f"     {kb:,.1f} KB")
